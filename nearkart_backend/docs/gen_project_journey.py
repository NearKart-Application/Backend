"""
NearKart — Complete Project Journey Document
Covers Sprint 0 to Sprint 12 — everything built and decided.
Run: python3 /tmp/gen_project_journey.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── PAGE MARGINS ──
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(0.9)
section.top_margin  = section.bottom_margin = Inches(0.9)

# ── HELPERS ──
def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def shade_para(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def h1(text):
    p = doc.add_heading(text, level=1)
    return p

def h2(text):
    return doc.add_heading(text, level=2)

def h3(text):
    return doc.add_heading(text, level=3)

def para(text, bold=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def code(text, indent=0.2):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    shade_para(p, 'F0F0F0')
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x8C)
    return p

def note(text, color='FFF3CD', text_color=(0x7D, 0x4E, 0x00)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    shade_para(p, color)
    run = p.add_run('  ' + text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(*text_color)
    return p

def success(text):
    return note('✅  ' + text, 'D4EDDA', (0x15, 0x57, 0x24))

def warning(text):
    return note('⚠️  ' + text, 'FFF3CD', (0x7D, 0x4E, 0x00))

def tip(text):
    return note('💡  ' + text, 'D1ECF1', (0x0C, 0x54, 0x60))

def decision(text):
    return note('🏗️  DESIGN DECISION: ' + text, 'EDE7F6', (0x4A, 0x14, 0x8C))

def sprint_box(number, title, status='Done ✅'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    shade_para(p, '1F4E79')
    r1 = p.add_run(f'  Sprint {number}')
    r1.font.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = RGBColor(0xFF, 0xD7, 0x00)
    r2 = p.add_run(f'  —  {title}')
    r2.font.bold = True
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r3 = p.add_run(f'   [{status}]')
    r3.font.size = Pt(10)
    r3.font.italic = True
    r3.font.color.rgb = RGBColor(0xCC, 0xFF, 0xCC)
    return p

def table(headers, rows, col_widths=None, header_color='1F4E79', alt_color='DEEAF1'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, header_color)
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            if ri % 2 == 0:
                shade_cell(cell, alt_color)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def divider():
    p = doc.add_paragraph('─' * 90)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.runs[0]
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('NearKart Backend')
run.font.name = 'Calibri'
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Complete Project Journey')
run2.font.size = Pt(20)
run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("India's First Hyperlocal Video Commerce Platform")
run3.font.size = Pt(13)
run3.font.italic = True
run3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run('Sprint 0 → Sprint 12   |   Full Backend Built   |   May 2026')
run4.font.size = Pt(11)
run4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

note(
    'This document covers the COMPLETE NearKart backend project — '
    'every sprint, every design decision, every model, every API endpoint, '
    'and every key technical choice made from the very first setup to the fully production-ready system.\n\n'
    'It is structured as a project journal: what was built, why it was built that way, '
    'and what the outcome was. Read it sprint-by-sprint or jump to any section.',
    'D1ECF1', (0x0C, 0x54, 0x60)
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 1 — WHAT IS NEARKART
# ════════════════════════════════════════════════════════════════
h1('1. What Is NearKart?')

para(
    'NearKart is India\'s first hyperlocal video commerce platform. It connects customers '
    'with nearby physical stores through short product videos — like TikTok shop meets '
    'Google Maps, but built specifically for Indian neighbourhood retail.',
    size=11
)

doc.add_paragraph()
h2('1.1  The Problem It Solves')
bullet('Local vendors (fashion, jewellery, footwear, electronics, food) have no digital presence')
bullet('Customers cannot discover nearby stores or see their current inventory')
bullet('No platform bridges the gap between "walk into a store" and "browse online"')
bullet('Small vendors cannot afford expensive e-commerce setups')

h2('1.2  How NearKart Solves It')
bullet('Vendors upload short videos (≤60s) of their products — no studio needed, just a phone')
bullet('Customers open the app → see a video feed of nearby stores within 1–5 km')
bullet('Customers can: follow stores, reserve products for pickup, chat live with vendors, join shopping groups')
bullet('Vendors get: video-based product showcase, chat inbox, reservation management, wallet + subscription')

h2('1.3  Full Platform Architecture')
code(
    'NEARKART PLATFORM\n'
    '┌─────────────────────────────────────────────────────────┐\n'
    '│  Customer App (React Native)  │  Vendor App (React Native)  │  Vendor Web Dashboard  │\n'
    '└────────────────────────────────────────────────────────────┘\n'
    '                            │  REST API + WebSocket\n'
    '                            ▼\n'
    '┌─────────────────────────────────────────────────────────────┐\n'
    '│                    BACKEND (this project)                    │\n'
    '│                                                              │\n'
    '│   Django REST API  •  Channels WebSocket  •  Celery Tasks    │\n'
    '│   PostgreSQL + PostGIS  •  Redis  •  AWS S3                 │\n'
    '│   Twilio SMS  •  Firebase FCM  •  Razorpay Payments         │\n'
    '│   Sentry monitoring  •  Nginx  •  AWS ECS (production)      │\n'
    '└─────────────────────────────────────────────────────────────┘'
)

h2('1.4  Technology Stack')
table(
    ['Layer', 'Technology', 'Why Chosen'],
    [
        ('Web Framework',      'Django 4.2 + Django REST Framework', 'Batteries included, ORM, admin panel, migrations'),
        ('Database',           'PostgreSQL 15 + PostGIS',            'PostGIS enables geo queries (nearby stores/products/videos)'),
        ('Cache',              'Redis (db=1)',                        'Fast cache for store detail, nearby results'),
        ('Task Queue',         'Celery + Redis broker (db=0)',        'Background transcoding, expired reservations, notifications'),
        ('WebSocket',          'Django Channels + Redis (db=2)',       'Real-time chat between customer and vendor'),
        ('ASGI Server',        'Daphne → Gunicorn + Uvicorn (prod)',   'Daphne for dev, Uvicorn workers for async performance in prod'),
        ('File Storage',       'AWS S3 (videos, images, static)',      'Scalable, presigned URLs so Django never handles binary data'),
        ('Video Transcoding',  'FFmpeg + HLS segments on S3',         'HLS works across all devices, segments allow adaptive streaming'),
        ('Auth',               'OTP via Twilio + JWT tokens',          'Passwordless login — friction-free for Indian mobile users'),
        ('Push Notifications', 'Firebase FCM',                        'Cross-platform (iOS + Android) push support'),
        ('Payments',           'Razorpay',                            'India-first, supports UPI/cards/wallets, strong webhook support'),
        ('Error Monitoring',   'Sentry',                              'Real-time error tracking in staging + production'),
        ('API Docs',           'drf-spectacular (Swagger UI)',         'Auto-generated from code, always in sync with endpoints'),
        ('Containerisation',   'Docker + Docker Compose',             'One command setup, identical dev/staging/prod environments'),
        ('CI/CD',              'GitHub Actions → AWS ECR → ECS',       'Automated lint, test, build, deploy with staging gate'),
    ],
    col_widths=[1.4, 2.0, 3.2],
)

h2('1.5  Sprint Roadmap (Overview)')
table(
    ['Sprint', 'Module', 'Key Deliverable', 'Status'],
    [
        ('S0',  'Environment',          'All tools, accounts, Docker, venv ready',           'Done ✅'),
        ('S1',  'Django Foundation',    'Running stack: Django + PostgreSQL + Redis + Celery', 'Done ✅'),
        ('S2',  'Auth Module',          'OTP login, JWT tokens, User model, roles',           'Done ✅'),
        ('S3',  'Store + Product',      'Stores, products, geo-search, reviews, follows',     'Done ✅'),
        ('S4',  'Video Module',         'S3 upload, HLS transcoding, location-based feed',    'Done ✅'),
        ('S5',  'Chat (WebSocket)',      'Real-time customer↔vendor chat + offline push',      'Done ✅'),
        ('S6',  'Blacklist Engine',      'Per-store customer blocking across all interactions','Done ✅'),
        ('S7',  'Billing + Wallet',      'Wallet, plans (Free/Basic/Premium), subscriptions',  'Done ✅'),
        ('S8',  'Analytics + Admin',    'Vendor dashboard, platform admin panel, stats',      'Done ✅'),
        ('S9',  'Reservations',          '2-hour product hold, confirm/cancel/expire flow',    'Done ✅'),
        ('S10', 'Groups',               'Collaborative shopping groups, Profile ID system',   'Done ✅'),
        ('S11', 'Notifications',        '18 notification types, FCM push + in-app inbox',     'Done ✅'),
        ('S12', 'Production + Razorpay','Production deploy, Razorpay payment flow',           'Done ✅'),
    ],
    col_widths=[0.5, 1.6, 3.4, 1.1],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 0
# ════════════════════════════════════════════════════════════════
sprint_box(0, 'Environment Setup')
doc.add_paragraph()

h2('What Was Set Up')
table(
    ['Tool / Service', 'Version / Detail', 'Purpose'],
    [
        ('Python',          '3.14.x',           'Backend language'),
        ('Docker Desktop',  '29.x',             'Run all services in containers'),
        ('Docker Compose',  'v5.x',             'Orchestrate 7 services with one command'),
        ('Git',             '2.x',              'Version control'),
        ('VS Code',         'latest',           'IDE'),
        ('Virtual env',     'venv/ in project', 'Isolated Python packages'),
        ('GitHub',          'repo created',     'Code repository for all sprints'),
    ],
    col_widths=[1.5, 1.5, 3.6],
)

h2('External Accounts Required')
table(
    ['Service', 'Purpose', 'Sprint First Needed'],
    [
        ('GitHub',       'Code repository',                  'All sprints'),
        ('AWS',          'S3 storage, ECS hosting',          'S4 (video), S12 (deploy)'),
        ('Firebase',     'FCM push notifications',           'S5 (chat), S11 (notifications)'),
        ('Twilio',       'SMS OTP delivery',                 'S2 (auth)'),
        ('Razorpay',     'Payment gateway',                  'S12 (billing)'),
        ('Sentry',       'Error monitoring',                 'S12 (production)'),
        ('Google Cloud', 'Maps reverse geocoding',           'S3 (store locality)'),
    ],
    col_widths=[1.2, 2.5, 2.0],
)

tip('For Sprints 1–11: only GitHub is strictly required. All third-party services have dev-mode bypass '
    '(placeholder credentials detected → mock responses returned, nothing real called).')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 1
# ════════════════════════════════════════════════════════════════
sprint_box(1, 'Django Foundation')
doc.add_paragraph()

para('Set up the complete Django project structure inside Docker. After this sprint the team '
     'can run the entire stack with one command and see Swagger UI + Django Admin in the browser.', size=11)

h2('Docker Services (7 containers)')
table(
    ['Service', 'Container Port', 'Docker Image', 'Purpose'],
    [
        ('Django REST API',     '8000', 'python:3.13-slim (custom)', 'Handles all HTTP requests'),
        ('Daphne (WebSocket)',   '8001', 'python:3.13-slim (custom)', 'WebSocket ASGI server'),
        ('Celery Worker',        '—',    'python:3.13-slim (custom)', 'Background tasks (video, SMS)'),
        ('Celery Beat',          '—',    'python:3.13-slim (custom)', 'Scheduled tasks (cron)'),
        ('PostgreSQL + PostGIS', '5432', 'postgis/postgis:15-3.3',    'Main database + geo extensions'),
        ('Redis',                '6379', 'redis:7-alpine',             'Cache, Celery broker, Channels layer'),
        ('Nginx',                '80',   'nginx:alpine',               'Rate limiting, proxy, static files'),
    ],
    col_widths=[1.8, 1.0, 2.2, 2.0],
)

h2('Project Folder Structure')
code(
    'nearkart_backend/\n'
    '├── config/\n'
    '│   ├── settings/\n'
    '│   │   ├── base.py          ← All shared settings (DB, JWT, Celery, AWS, Razorpay)\n'
    '│   │   ├── development.py   ← Local dev: DEBUG=True, dev OTP\n'
    '│   │   ├── staging.py       ← Staging: inherits production, Swagger allowed\n'
    '│   │   ├── production.py    ← Production: S3, HSTS, CloudWatch logging\n'
    '│   │   └── testing.py       ← Testing: SQLite, no Docker needed\n'
    '│   ├── asgi.py              ← WebSocket entry point\n'
    '│   ├── celery.py            ← Celery app + autodiscover tasks\n'
    '│   ├── urls.py              ← All API routes wired here\n'
    '│   └── wsgi.py              ← HTTP entry point\n'
    '│\n'
    '├── core/\n'
    '│   ├── models.py            ← BaseModel (UUID pk + timestamps)\n'
    '│   ├── exceptions.py        ← Custom error format: {error, message, code, details}\n'
    '│   ├── permissions.py       ← IsCustomer, IsVendor, IsAdmin, IsStoreOwner\n'
    '│   ├── pagination.py        ← StandardOffsetPagination, CursorPagination\n'
    '│   ├── middleware.py        ← JWT auth middleware for WebSocket connections\n'
    '│   └── utils/\n'
    '│       ├── cache.py         ← Redis cache helpers + TTL constants\n'
    '│       ├── geo.py           ← PostGIS DWithin geo-query helpers\n'
    '│       └── s3.py            ← AWS S3 presigned URL helpers\n'
    '│\n'
    '├── apps/                    ← All 13 Django apps\n'
    '│   ├── auth_app/\n'
    '│   ├── stores/\n'
    '│   ├── products/\n'
    '│   ├── videos/\n'
    '│   ├── chat/\n'
    '│   ├── billing/\n'
    '│   ├── analytics/\n'
    '│   ├── blacklist/\n'
    '│   ├── notifications/\n'
    '│   ├── reservations/\n'
    '│   ├── groups/\n'
    '│   └── admin_panel/\n'
    '│\n'
    '├── docker-compose.yml         ← Dev stack\n'
    '├── docker-compose.prod.yml    ← Production stack (no code mounts)\n'
    '├── Dockerfile                 ← Multi-stage: builder → development → production\n'
    '├── nginx/\n'
    '│   ├── nginx.conf             ← Dev nginx config\n'
    '│   └── nginx.prod.conf        ← Prod nginx (Swagger blocked, attack paths blocked)\n'
    '├── scripts/entrypoint.sh      ← Auto-migrate + collectstatic before gunicorn\n'
    '├── requirements/\n'
    '│   ├── base.txt              ← Core packages\n'
    '│   ├── development.txt       ← + testing/linting tools\n'
    '│   └── production.txt        ← + gunicorn/uvicorn\n'
    '└── .env.example               ← All env vars documented with comments'
)

h2('Key Configuration Decisions Made in Sprint 1')
table(
    ['Decision', 'Choice Made', 'Reason'],
    [
        ('Primary Key type',     'UUID everywhere (not integer)',        'UUIDs are non-sequential — safe to expose in URLs, no enumeration attacks'),
        ('Timezone',             'Asia/Kolkata (USE_TZ=True)',           'All users are in India; store datetimes in UTC, display in IST'),
        ('Error format',         'Custom handler: {error, message, details}', 'Consistent format across all 79 endpoints — frontend can always parse the same shape'),
        ('Redis databases',      '0=Celery broker, 1=cache, 2=Channels', 'Separate DBs prevent cache flushes from killing message queue'),
        ('JWT access token',     '1 hour lifetime',                     'Short enough to be secure; refresh token (30 days) handles re-auth transparently'),
        ('Settings split',       'base/development/staging/production',  'Env-specific settings without duplication; staging inherits production with relaxations'),
        ('BaseModel',            'All models inherit UUID pk + timestamps','Consistency across all 20+ models; never forget created_at/updated_at again'),
    ],
    col_widths=[1.6, 2.0, 3.0],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 2
# ════════════════════════════════════════════════════════════════
sprint_box(2, 'Auth Module — OTP Login + JWT')
doc.add_paragraph()

para('Passwordless authentication using Indian mobile numbers. Users send their phone number, '
     'receive a 6-digit OTP via SMS (Twilio), verify it, and receive JWT tokens. '
     'No passwords ever stored or transmitted.', size=11)

h2('Database Models')
table(
    ['Model', 'Key Fields', 'Purpose'],
    [
        ('User',        'phone_number (unique, +91XXXXXXXXXX), role (customer/vendor/admin), full_name, profile_id, registered_location (PointField)', 'Single user model for all roles'),
        ('OTPToken',    'user, otp_hash (SHA256), expires_at (10 min), is_used, attempts (locked at 5)', 'Secure OTP session — hash stored never plain text'),
        ('DeviceToken', 'user, fcm_token, device_type (android/ios/web), is_active', 'Firebase push token per device'),
    ],
    col_widths=[1.2, 3.0, 2.4],
)

h2('OTP Login Flow')
code(
    'Phone number entered by user\n'
    '        ↓\n'
    'POST /auth/otp/send/\n'
    '  → User created if first login (role auto-set: first digit 9=vendor, others=customer for dev)\n'
    '  → Random 6-digit OTP generated\n'
    '  → SHA256(OTP) stored in OTPToken (never plain text)\n'
    '  → Celery task queued → Twilio sends SMS\n'
    '  → Dev mode: OTP is always 123456 (DEV_FIXED_OTP in .env)\n'
    '        ↓\n'
    'POST /auth/otp/verify/\n'
    '  → SHA256(submitted OTP) compared with stored hash\n'
    '  → Checks: not expired, not used, attempts < 5\n'
    '  → Success → JWT access token (1h) + refresh token (30d) returned\n'
    '  → OTPToken marked is_used=True\n'
    '        ↓\n'
    'All subsequent requests use: Authorization: Bearer <access_token>'
)

h2('API Endpoints')
table(
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ('POST', '/auth/otp/send/',      'None', 'Send OTP to phone. Creates user if first login.'),
        ('POST', '/auth/otp/verify/',    'None', 'Verify OTP → return access + refresh tokens'),
        ('POST', '/auth/token/refresh/', 'None', 'Get new access token using refresh token'),
        ('GET',  '/auth/me/',            'JWT',  'Get current user profile + profile_id'),
        ('PATCH','/auth/me/',            'JWT',  'Update full_name, email'),
        ('PUT',  '/auth/me/location/',   'JWT',  'Update user\'s registered location (lat/lng)'),
        ('POST', '/auth/logout/',        'JWT',  'Blacklist refresh token — logout'),
        ('GET',  '/auth/users/search/',  'JWT',  'Search user by Profile ID (added Sprint 10)'),
    ],
    col_widths=[0.6, 2.0, 0.8, 3.2],
)

h2('Key Decisions')
decision('Phone number format enforced as +91XXXXXXXXXX with regex — first digit must be 6/7/8/9 (Indian mobile range). '
         'Rejects: 9999999999 (missing +91), 919999999999 (missing +), +91 9999... (space), +915000000000 (starts with 5).')
decision('OTP stored as SHA256 hash — even if DB is compromised, OTPs cannot be reversed. '
         'Plain OTP is never logged anywhere.')
decision('JWT refresh token blacklisted on logout using djangorestframework-simplejwt blacklist app. '
         'Prevents token reuse after logout.')
decision('Dev mode fixed OTP (123456) — set in .env so QA and development never need real Twilio integration. '
         'Production removes DEV_FIXED_OTP.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 3
# ════════════════════════════════════════════════════════════════
sprint_box(3, 'Store + Product Module')
doc.add_paragraph()

para('Vendors create stores with geographic location. Customers discover nearby stores and products '
     'using PostGIS spatial queries. Stores have hours, reviews, followers, and QR codes.', size=11)

h2('Store Model — Key Fields')
table(
    ['Field', 'Type', 'Detail'],
    [
        ('owner',            'OneToOneField → User',    'One store per vendor account'),
        ('location',         'PointField(geography=True)','PostGIS geographic coordinates — enables DWithin distance queries'),
        ('locality',         'CharField',               'Auto-filled by Google Maps reverse geocode on store create/location change'),
        ('category',         'Enum',                    'fashion/jewellery/footwear/decor/furniture/gifts/beauty/food/electronics/other'),
        ('is_verified',      'Boolean (default False)', 'Must be True to appear in nearby results — admin sets this'),
        ('is_open',          'Boolean (default False)', 'Vendor toggles; fires store_opened notifications to followers'),
        ('performance_score','Float',                   'Auto-calculated average of all review ratings'),
        ('wallet_balance',   'Decimal',                 'Vendor wallet — topped up via Razorpay, deducted on subscription'),
        ('qr_code_url',      'URLField',                'Generated by QRService on first GET /stores/<id>/qr-code/; uploaded to S3'),
    ],
    col_widths=[1.5, 1.8, 3.3],
)

h2('Product Model — Key Fields')
table(
    ['Field', 'Type', 'Detail'],
    [
        ('status',      'Enum: draft/active/inactive/out_of_stock', 'Only "active" products appear in nearby/search'),
        ('is_visible',  'Boolean',                                   'False = hidden from customers even if active'),
        ('base_price',  'Decimal',                                   'Reference price; actual price can be on variants'),
        ('ProductVariant', 'ForeignKey → Product',                   'name, sku (globally unique), price, stock_quantity'),
        ('ProductImage',   'ForeignKey → Product',                   'image_url, s3_key, is_primary, order'),
        ('Wishlist',       'user + product (unique_together)',        'Toggle — same endpoint adds or removes'),
    ],
    col_widths=[1.5, 2.0, 3.1],
)

h2('Store Endpoints')
table(
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ('GET',  '/stores/nearby/',            'None',       'DWithin PostGIS query — ?lat=&lng=&radius=&category='),
        ('GET',  '/stores/<uuid>/',            'None',       'Full detail — includes hours, reviews (cached in Redis)'),
        ('POST', '/stores/',                   'Vendor JWT', 'Create store (1 per vendor)'),
        ('PUT',  '/stores/<uuid>/update/',     'Owner JWT',  'Update store fields (partial)'),
        ('POST', '/stores/<uuid>/follow/',     'JWT',        'Follow/unfollow toggle'),
        ('POST', '/stores/<uuid>/review/',     'JWT',        'Add or update review (1–5 stars)'),
        ('GET',  '/stores/<uuid>/qr-code/',    'Owner JWT',  'Get or generate QR code'),
        ('GET',  '/stores/<uuid>/hours/',      'Owner JWT',  'Get operating hours'),
        ('PUT',  '/stores/<uuid>/hours/',      'Owner JWT',  'Set hours — replaces ALL existing hours atomically'),
    ],
    col_widths=[0.6, 2.2, 1.0, 2.8],
)

h2('Product Endpoints')
table(
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ('GET',    '/products/nearby/',         'None',       'Geo-search products within radius'),
        ('GET',    '/products/search/',         'None',       'Trigram similarity search by name (?q=)'),
        ('GET',    '/products/<uuid>/',         'None',       'Full detail with variants, images, is_wishlisted'),
        ('POST',   '/products/',                'Vendor JWT', 'Create product (checks plan video/product limits)'),
        ('PUT',    '/products/<uuid>/update/',  'Owner JWT',  'Update product (partial)'),
        ('DELETE', '/products/<uuid>/update/',  'Owner JWT',  'Hard delete product'),
        ('POST',   '/products/<uuid>/wishlist/','JWT',        'Add/remove wishlist toggle'),
    ],
    col_widths=[0.6, 2.2, 1.0, 2.8],
)

h2('Key Decisions')
decision('geography=True on Store.location — required for DWithin(D(km=N)) distance queries. '
         'Plain geometry field raises ValueError with D(km=...) units.')
decision('is_verified=False by default — new stores are invisible until admin verifies them. '
         'Prevents fake stores appearing in the customer feed immediately.')
decision('Trigram search threshold 0.2 — TrigramSimilarity > 0.2 catches partial name matches '
         'while filtering noise. Results sorted by similarity score descending.')
decision('Store Hours uses atomic replace-all on PUT — simpler than per-day PATCH. '
         'Client sends the complete week schedule; backend deletes old entries and bulk-creates new ones.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 4
# ════════════════════════════════════════════════════════════════
sprint_box(4, 'Video Module — Upload, Transcode, Feed')
doc.add_paragraph()

para('Vendors upload product videos directly to S3 (Django never handles the binary). '
     'Celery + FFmpeg transcodes to HLS for adaptive streaming. '
     'Customers see a location-based video feed sorted by distance.', size=11)

h2('Two-Step Upload Flow')
code(
    'Step 1 — Request upload URL\n'
    'POST /api/v1/videos/request-upload/\n'
    '  → Django calls S3 to generate presigned PUT URL (expires 15 min)\n'
    '  → Returns: {video_id, upload_url, s3_key}\n'
    '  → Video record created with status="pending_upload"\n'
    '\n'
    'Step 2 — Vendor uploads directly to S3 (NOT through Django)\n'
    'PUT <upload_url> (with video file binary)\n'
    '  → S3 stores raw MP4\n'
    '  → Django server is NOT involved — no memory pressure\n'
    '\n'
    'Step 3 — Confirm upload\n'
    'POST /api/v1/videos/<id>/confirm-upload/\n'
    '  → Validates duration_seconds ≤ 60 (VIDEO_MAX_DURATION_SECONDS)\n'
    '  → Changes status to "processing"\n'
    '  → Queues transcode_video Celery task\n'
    '\n'
    'Celery Task: transcode_video\n'
    '  → Downloads raw MP4 from S3\n'
    '  → FFmpeg: convert to HLS (.m3u8 + .ts segments)\n'
    '  → Upload HLS files back to S3\n'
    '  → Set status="ready", hls_url set, thumbnail generated\n'
    '  → notify_video_ready notification sent to vendor\n'
    '  → Dev mode: FFmpeg skipped, video immediately marked ready'
)

h2('Video Model — Key Fields')
table(
    ['Field', 'Type', 'Purpose'],
    [
        ('store',          'ForeignKey → Store',    'Whose store this video belongs to'),
        ('location',       'PointField',             'Copied from store — feeds use this for geo index, not store.location'),
        ('status',         'Enum',                   'pending_upload → processing → ready / failed'),
        ('raw_s3_key',     'CharField',              'S3 key of original MP4'),
        ('hls_url',        'URLField',               'CDN URL of HLS .m3u8 playlist — set after transcoding'),
        ('thumbnail_url',  'URLField',               'First frame extracted by FFmpeg'),
        ('duration_seconds','IntegerField',           'Validated ≤ 60s on confirm-upload'),
        ('view_count',     'PositiveIntegerField',    'Incremented atomically via F() on each GET /videos/<id>/'),
        ('like_count',     'PositiveIntegerField',    'Incremented/decremented via F() on like toggle'),
        ('expires_at',     'DateTimeField',           '30 days from upload; daily Celery task deletes expired videos'),
    ],
    col_widths=[1.5, 1.5, 3.6],
)

h2('Endpoints')
table(
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ('POST', '/videos/request-upload/',       'Vendor JWT', 'Get presigned S3 URL + video_id'),
        ('POST', '/videos/<id>/confirm-upload/',  'Vendor JWT', 'Trigger Celery transcoding'),
        ('GET',  '/videos/my-videos/',            'Vendor JWT', 'Vendor library (all statuses, ?status= filter)'),
        ('GET',  '/videos/feed/',                 'None',       'Location-based feed (?lat=&lng=&radius=)'),
        ('GET',  '/videos/<id>/',                 'None',       'Detail + view count increment'),
        ('PATCH','/videos/<id>/update/',          'Vendor JWT', 'Update title/description/is_visible'),
        ('DELETE','/videos/<id>/delete/',         'Vendor JWT', 'Permanent delete (own videos only)'),
        ('POST', '/videos/<id>/like/',            'JWT',        'Like/unlike toggle + video_liked notification'),
    ],
    col_widths=[0.6, 2.2, 1.0, 2.8],
)

h2('Key Decisions')
decision('Presigned S3 URL — Django never receives the video binary. For a 100MB video, '
         'Django would need 100MB RAM per upload. With presigned URLs, vendor uploads directly to S3 '
         'and Django just creates a DB record.')
decision('HLS over MP4 — HLS (.m3u8 + .ts segments) enables adaptive bitrate streaming. '
         'Devices can switch quality mid-play based on network speed. Works natively in browsers and iOS/Android.')
decision('Video location copied from store at create time — feed geo queries use a spatial index on '
         'Video.location, not a JOIN to stores. JOIN with geo filter would be slow at scale.')
decision('F() expressions for view_count and like_count — prevents race condition where two '
         'concurrent requests both read count=5, both write count=6. F() does atomic SQL: UPDATE SET count = count + 1.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 5
# ════════════════════════════════════════════════════════════════
sprint_box(5, 'Chat Module — Real-Time WebSocket')
doc.add_paragraph()

para('Real-time bidirectional chat between customer and vendor. When online, messages delivered '
     'instantly via WebSocket. When offline, FCM push notification is sent. One conversation per '
     'customer-store pair.', size=11)

h2('Architecture')
code(
    'Customer app                    Vendor app\n'
    '    │                               │\n'
    '    └─── WebSocket ────────────────┘\n'
    '              │\n'
    '         Django Channels\n'
    '              │\n'
    '         Redis Channel Layer\n'
    '              │\n'
    '         ChatConsumer (async)\n'
    '              │\n'
    '         ┌────┴────┐\n'
    '         │         │\n'
    '     Online:    Offline:\n'
    '  WS broadcast  FCM push\n'
    '     to group   via Firebase'
)

h2('Models')
table(
    ['Model', 'Key Fields', 'Purpose'],
    [
        ('Conversation', 'customer, store, unread_count_customer, unread_count_vendor, last_message_at (indexed)', 'One per (customer, store) pair — unique_together'),
        ('Message', 'conversation, sender, content, message_type (text/image/product_ref/video_ref), is_read', 'Individual messages; supports future media sharing'),
    ],
    col_widths=[1.2, 3.2, 2.2],
)

h2('WebSocket Connection')
code(
    'URL: ws://localhost:8001/ws/conversations/<conversation_id>/?token=<jwt_access_token>\n'
    '\n'
    'Connection flow:\n'
    '  1. JWTAuthMiddleware reads ?token= query param → authenticates user\n'
    '  2. Consumer verifies user is part of the conversation → else close 4003\n'
    '  3. Consumer joins Redis channel group: "conversation_<id>"\n'
    '  4. Both participants in the group receive all messages in real-time\n'
    '\n'
    'Message format (client sends):\n'
    '  {"type": "chat_message", "content": "Hello!"}\n'
    '\n'
    'Message format (server broadcasts to group):\n'
    '  {"type": "chat_message", "id": "...", "sender_id": "...", "content": "Hello!", "created_at": "..."}\n'
    '\n'
    'Error close codes:\n'
    '  4001 — invalid or missing JWT token\n'
    '  4003 — conversation not found or user not a participant (also: blacklisted)'
)

h2('REST Endpoints')
table(
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ('POST', '/conversations/start/',              'JWT', 'Get-or-create conversation (returns 200 if exists, 201 if new)'),
        ('GET',  '/conversations/',                    'JWT', 'List all conversations — inbox sorted by last_message_at desc'),
        ('GET',  '/conversations/<id>/messages/',      'JWT', 'Paginated history (50/page, ?before=<message_id> for older)'),
        ('PATCH','/conversations/<id>/read/',          'JWT', 'Reset my unread count to 0'),
    ],
    col_widths=[0.6, 2.4, 0.6, 3.0],
)

h2('Key Decisions')
decision('One conversation per (customer, store) unique_together — no duplicate threads. '
         'start/ is idempotent: returns existing conversation if it exists.')
decision('JWT via ?token= WebSocket query param — WebSocket connections cannot send custom HTTP headers. '
         'The existing JWTAuthMiddleware intercepts it before the consumer is instantiated.')
decision('Unread counts split: unread_count_customer + unread_count_vendor — allows showing separate '
         'badge counts in each party\'s inbox without a query.')
decision('message_type enum (text/image/product_ref/video_ref) — extensible for future '
         'in-chat product sharing without schema changes.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 6
# ════════════════════════════════════════════════════════════════
sprint_box(6, 'Blacklist Engine')
doc.add_paragraph()

para('Vendors can block specific customers from all interactions with their store. '
     'Blocking is per-store (not per-vendor account), toggled on a single endpoint, '
     'and enforced at every interaction point.', size=11)

h2('Enforcement Matrix')
table(
    ['Interaction', 'What Blocked Customer Gets'],
    [
        ('POST /stores/<id>/follow/',      '403 — blacklisted'),
        ('POST /stores/<id>/review/',      '403 — blacklisted'),
        ('POST /conversations/start/',     '403 — blacklisted'),
        ('WebSocket connect',              'Close code 4003 (same as "not a member" — no info leak)'),
        ('POST /reservations/',            '403 — blacklisted'),
    ],
    col_widths=[2.8, 3.8],
)

h2('Endpoints')
table(
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ('POST', '/stores/<store_id>/blacklist/<customer_id>/', 'Vendor JWT', 'Block if not blocked; unblock if blocked (toggle)'),
        ('GET',  '/stores/<store_id>/blacklist/',               'Vendor JWT', 'List all blocked customers with reason + date'),
    ],
    col_widths=[0.6, 3.0, 1.0, 2.0],
)

h2('Key Decisions')
decision('Per-store not per-vendor — a vendor could manage multiple stores in future. '
         'Each store has its own independent blacklist.')
decision('Toggle on single endpoint — POST on same URL alternates between block and unblock. '
         'No separate DELETE endpoint — simpler REST surface.')
decision('WS close code 4003 (not a specific "blocked" code) — prevents the blocked customer '
         'from knowing they are specifically blocked vs just not a participant.')
decision('Read history preserved — blocked customer can still read old messages in '
         'GET /conversations/<id>/messages/ but cannot send new messages.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 7
# ════════════════════════════════════════════════════════════════
sprint_box(7, 'Billing + Wallet + Subscription Plans')
doc.add_paragraph()

para('Vendors have a wallet (Decimal balance on the Store model). They top up the wallet and '
     'use the balance to subscribe to a plan. The plan controls how many videos and products '
     'they can have active simultaneously.', size=11)

h2('Subscription Plans')
table(
    ['Plan', 'Monthly Price', 'Active Videos', 'Active Products', 'Target Vendor'],
    [
        ('Free',    '₹0',    '3',         '10',        'New vendors getting started'),
        ('Basic',   '₹499',  '20',        '50',        'Growing stores with regular uploads'),
        ('Premium', '₹999',  'Unlimited', 'Unlimited', 'Large stores with high inventory'),
    ],
    col_widths=[0.9, 1.2, 1.2, 1.4, 2.9],
)

h2('Billing Models')
table(
    ['Model', 'Key Fields', 'Purpose'],
    [
        ('Plan',        'name (slug), display_name, price, duration_days, video_limit, product_limit', 'Master plan config — seeded once via manage.py seed_plans'),
        ('Subscription','store (OneToOne), plan, started_at, expires_at, is_active', 'One active subscription per store — updated in-place on renewal'),
        ('Transaction', 'store, type (topup/subscription/refund), amount, reference_id, balance_after', 'Full audit trail of all wallet movements'),
    ],
    col_widths=[1.2, 3.0, 2.4],
)

h2('Plan Limit Enforcement')
table(
    ['Action', 'Check', 'Response When Over Limit'],
    [
        ('POST /videos/request-upload/', 'Count ready+processing+pending videos vs plan.video_limit', '403 — plan_limit_reached'),
        ('POST /products/',              'Count all non-deleted products vs plan.product_limit',       '403 — plan_limit_reached'),
    ],
    col_widths=[2.2, 2.8, 1.8],
)

h2('Celery Beat Task')
bullet('Task: billing.expire_subscriptions')
bullet('Schedule: Daily at midnight IST')
bullet('Action: Marks all Subscription records with expires_at < now() as is_active=False')
bullet('Effect: Expired vendors fall back to Free plan limits on next upload/create')

h2('Key Decisions')
decision('Wallet on Store model (not a separate Wallet model) — keeps the data model simple. '
         'Store.wallet_balance is the single source of truth. All changes go through atomic '
         'select-for-update to prevent race conditions on concurrent top-ups.')
decision('Subscription update_or_create in-place — one row per store, not an append-only history. '
         'Transaction table provides the audit trail. This keeps subscription queries O(1).')
decision('Free plan costs ₹0 — subscribe endpoint accepts free plan without touching wallet balance. '
         'Allows testing the full subscribe flow without needing real money.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 8
# ════════════════════════════════════════════════════════════════
sprint_box(8, 'Analytics + Admin Panel')
doc.add_paragraph()

para('Two separate read-heavy modules with no new database models — all data aggregated '
     'from existing tables. Analytics gives vendors their store performance dashboard. '
     'Admin panel gives staff full platform visibility and control.', size=11)

h2('Vendor Analytics Dashboard — Response Shape')
code(
    '{\n'
    '  "store": { "name", "category", "is_active", "is_verified", "is_open",\n'
    '             "follower_count", "review_count", "avg_rating" },\n'
    '  "wallet": { "balance" },\n'
    '  "subscription": { "plan", "expires_at", "is_active", "days_left" },\n'
    '  "current_plan": { "name", "display_name", "video_limit", "product_limit" },\n'
    '  "products": { "total", "active", "draft", "inactive" },\n'
    '  "videos":   { "total", "ready", "processing", "pending", "total_likes", "total_views" }\n'
    '}'
)

h2('Analytics Endpoints')
table(
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ('GET', '/analytics/vendor/',          'Vendor JWT', 'Full store performance snapshot'),
        ('GET', '/analytics/vendor/videos/',   'Vendor JWT', 'Per-video view + like counts'),
        ('GET', '/analytics/vendor/products/', 'Vendor JWT', 'Per-product wishlist counts'),
    ],
    col_widths=[0.6, 2.2, 1.0, 2.8],
)

h2('Admin Panel Endpoints')
table(
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ('GET',   '/admin-panel/stats/',                       'Staff JWT', 'Platform-wide: users, stores, videos, revenue'),
        ('GET',   '/admin-panel/stores/',                      'Staff JWT', 'All stores — filter by is_verified, category, search'),
        ('PATCH', '/admin-panel/stores/<id>/',                 'Staff JWT', 'Verify/activate/deactivate a store'),
        ('GET',   '/admin-panel/users/',                       'Staff JWT', 'All users — filter by role, is_active, search'),
        ('POST',  '/admin-panel/users/<id>/toggle-active/',    'Staff JWT', 'Enable or disable a user account'),
    ],
    col_widths=[0.6, 2.6, 0.8, 2.6],
)

h2('How to Create a Staff User (Dev)')
code(
    'docker compose exec django python manage.py shell -c "\n'
    'from apps.auth_app.models import User\n'
    'u, _ = User.objects.get_or_create(phone_number=\'+919000000001\', defaults={\'role\':\'admin\'})\n'
    'u.is_staff = True; u.is_superuser = True; u.save()\n'
    'print(\'Done\')\n'
    '"'
)
para('Then login: POST /auth/otp/send/ with +919000000001 → OTP 123456 → get admin JWT.', size=10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 9
# ════════════════════════════════════════════════════════════════
sprint_box(9, 'Reservations — 2-Hour Product Hold')
doc.add_paragraph()

para('Customers can hold a product at a nearby store for 2 hours. '
     'The vendor receives a notification, confirms the hold, and marks it complete when the '
     'customer arrives. Celery automatically expires unclaimed holds every hour.', size=11)

h2('Reservation State Machine')
code(
    '             ┌──────────────────────┐\n'
    '  Customer   │   CREATE             │\n'
    '  ─────────► │   pending            │\n'
    '             └──────────────────────┘\n'
    '                     │\n'
    '          ┌──────────┼──────────────────┐\n'
    '          ▼          ▼                  ▼\n'
    '      confirmed  cancelled (cust/vendor) expired (Celery)\n'
    '          │\n'
    '     ┌────┴────────────┐\n'
    '     ▼                 ▼\n'
    '  completed        cancelled (vendor)'
)

h2('Endpoints')
table(
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ('POST',  '/reservations/',              'JWT',        'Customer creates reservation (store + product + quantity + note)'),
        ('GET',   '/reservations/list/',         'JWT',        'Customer: own reservations. Vendor: store\'s reservations.'),
        ('GET',   '/reservations/<id>/',         'JWT',        'Reservation detail (customer or store vendor only)'),
        ('PATCH', '/reservations/<id>/status/',  'Vendor JWT', 'Confirm / cancel / complete'),
        ('POST',  '/reservations/<id>/cancel/',  'JWT',        'Customer cancels own pending reservation'),
    ],
    col_widths=[0.6, 2.0, 1.0, 3.0],
)

h2('Celery Beat Task')
bullet('Task: reservations.expire_reservations')
bullet('Schedule: Top of every hour (:00)')
bullet('Action: All pending reservations with expires_at < now() → status = expired')
bullet('Side effect: Sends reservation_expired notification to the customer')

h2('Key Decision')
decision('Hold duration as env var (RESERVATION_HOLD_HOURS=2) — different store categories '
         'might need different hold windows. Keeping it configurable avoids hardcoding.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 10
# ════════════════════════════════════════════════════════════════
sprint_box(10, 'Groups — Collaborative Shopping')
doc.add_paragraph()

para('Users create groups (like a WhatsApp group but for shopping). Members share products, '
     'discuss, and an admin finalizes the group\'s choice. Phone numbers are NEVER shared — '
     'users are identified by a private Profile ID (NK-XXXXXXXX format).', size=11)

h2('Profile ID System')
table(
    ['Aspect', 'Detail'],
    [
        ('Format',       'NK-XXXXXXXX (8 random alphanumeric chars, uppercase)'),
        ('Assignment',   'Auto-generated on user signup — stored in User.profile_id'),
        ('Visibility',   'Only shown in GET /auth/me/ — user must share it manually'),
        ('Search',       'GET /auth/users/search/?profile_id=NK-A3X9K2 → returns name only (NO phone number)'),
        ('Purpose',      'Add friends to groups without exposing phone numbers — privacy-first design'),
    ],
    col_widths=[1.3, 5.3],
)

h2('Two Group Types')
table(
    ['Type', 'Creator', 'Who Can Join', 'How Members Are Added'],
    [
        ('customer', 'Any user',  'Anyone (by Profile ID)',          'Admin sends {profile_id} → user looked up and added'),
        ('vendor',   'Vendor',    'Customers who follow that store', 'Admin fetches eligible-members list → adds by {user_id}'),
    ],
    col_widths=[0.8, 1.0, 2.0, 2.8],
)

h2('Group Flow')
code(
    'Create group → Add members → Share products → Finalize → Done\n'
    '\n'
    '1. Creator creates group (auto-becomes admin)\n'
    '2. Admin adds members: by profile_id (customer groups) or user_id from eligible-members (vendor groups)\n'
    '3. Any member shares a product: POST /groups/<id>/products/ {product_id, note}\n'
    '4. All other members are notified\n'
    '5. Admin marks one as final: POST /groups/<id>/products/<sp_id>/finalize/\n'
    '6. All members notified of final choice\n'
    '7. Creator can delete group / any member can leave'
)

h2('Business Rules')
table(
    ['Rule', 'Detail'],
    [
        ('No phone in responses',    'All group responses use profile_id + full_name only — zero phone exposure'),
        ('Creator protection',       'Creator cannot be removed, demoted, or leave — must delete the group'),
        ('Multiple admins',          'Any admin can promote/demote others (but not the creator)'),
        ('External link blocking',   'notes with http:// or https:// → 400 — External links not allowed. Only nearkart:// app links OK'),
        ('Vendor group restriction', 'Non-follower cannot be added → 403. Eligible-members list ensures only followers shown'),
        ('Product must be active',   'Shared product must have status=active + is_visible=True'),
    ],
    col_widths=[1.8, 4.8],
)

h2('Endpoints (14 total)')
table(
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ('GET',    '/auth/users/search/',                        'JWT',        'Search by Profile ID'),
        ('POST',   '/groups/',                                   'JWT',        'Create group (customer or vendor type)'),
        ('GET',    '/groups/',                                   'JWT',        'List my groups'),
        ('GET',    '/groups/<id>/',                              'Member JWT', 'Group detail + members list'),
        ('DELETE', '/groups/<id>/',                              'Creator JWT','Delete group'),
        ('POST',   '/groups/<id>/members/add/',                  'Admin JWT',  'Add member by profile_id or user_id'),
        ('DELETE', '/groups/<id>/members/<uid>/remove/',         'Admin JWT',  'Remove member'),
        ('POST',   '/groups/<id>/members/<uid>/make-admin/',     'Admin JWT',  'Promote to admin'),
        ('POST',   '/groups/<id>/members/<uid>/remove-admin/',   'Admin JWT',  'Demote from admin'),
        ('POST',   '/groups/<id>/leave/',                        'Member JWT', 'Leave group'),
        ('GET',    '/groups/<id>/eligible-members/',             'Admin JWT',  'Vendor: followers not yet in group'),
        ('GET',    '/groups/<id>/products/',                     'Member JWT', 'List shared products'),
        ('POST',   '/groups/<id>/products/',                     'Member JWT', 'Share a product in group'),
        ('POST',   '/groups/<id>/products/<sp_id>/finalize/',    'Admin JWT',  'Finalize the group\'s product choice'),
    ],
    col_widths=[0.6, 2.8, 0.8, 2.4],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 11
# ════════════════════════════════════════════════════════════════
sprint_box(11, 'Notifications — In-App Inbox + FCM Push')
doc.add_paragraph()

para('Every important event in the system sends two things simultaneously: '
     '(1) a Notification record saved to the database (in-app inbox), '
     '(2) a Firebase FCM push notification to the user\'s device(s).', size=11)

h2('Architecture')
code(
    'Any service method (billing, reservations, chat, groups, stores, videos)\n'
    '        ↓\n'
    'NotificationService.notify_*()  (one of 18 helper methods)\n'
    '        ↓\n'
    '    ┌───────────────────┐\n'
    '    │  Notification.objects.create()   ← In-app inbox (database)\n'
    '    │  FCMService.send_push()           ← Firebase → device\n'
    '    └───────────────────┘\n'
    '\n'
    'Dev mode (no Firebase credentials):\n'
    '  [FCM-DEV] → +919876543210 | title="New message" body="..." tokens=1\n'
    '  (logged to console — nothing sent to Firebase)'
)

h2('18 Notification Types and Their Triggers')
table(
    ['Notification Type', 'Triggered By', 'Recipient'],
    [
        ('new_message',             'ChatConsumer: customer/vendor sends WS message',          'Other party in conversation'),
        ('reservation_created',     'Customer creates reservation',                            'Vendor (store owner)'),
        ('reservation_confirmed',   'Vendor confirms reservation',                             'Customer'),
        ('reservation_cancelled',   'Vendor or customer cancels',                              'Customer'),
        ('reservation_expired',     'Celery: hold expired after 2 hours',                     'Customer'),
        ('new_follower',            'User follows a store',                                    'Vendor'),
        ('new_review',              'User submits a store review',                             'Vendor'),
        ('store_opened',            'Vendor sets is_open=true on a closed store',              'All followers (bulk)'),
        ('video_liked',             'User likes a video',                                      'Vendor'),
        ('video_ready',             'Celery: transcoding complete',                            'Vendor'),
        ('wallet_topup',            'Razorpay verify OR admin topup credits wallet',           'Vendor'),
        ('subscription_expiring',   'Celery: daily 9 AM — sub expires in ~3 days',            'Vendor'),
        ('subscription_expired',    'Celery: daily 9:05 AM — sub expired in last 24h',        'Vendor'),
        ('group_added',             'Admin adds member to group',                              'Added user'),
        ('group_removed',           'Admin removes member from group',                         'Removed user'),
        ('group_product_shared',    'Member shares product in group',                          'All other group members'),
        ('group_product_finalized', 'Admin finalizes group product',                           'All group members'),
        ('group_admin_promoted',    'Admin promotes member to admin',                          'Promoted user'),
    ],
    col_widths=[1.9, 2.6, 2.1],
)

h2('Inbox Endpoints')
table(
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ('POST', '/notifications/device-token/',  'JWT', 'Register or update FCM token'),
        ('GET',  '/notifications/',               'JWT', 'Last 50 notifications (inbox)'),
        ('GET',  '/notifications/unread-count/',  'JWT', 'Badge count for app icon'),
        ('POST', '/notifications/<id>/read/',     'JWT', 'Mark one read (idempotent)'),
        ('POST', '/notifications/read-all/',      'JWT', 'Mark all read'),
    ],
    col_widths=[0.6, 2.2, 0.8, 3.0],
)

h2('Celery Beat Tasks')
table(
    ['Task', 'Schedule', 'Action'],
    [
        ('notify_expiring_subscriptions', 'Daily 9:00 AM IST', 'Finds vendors whose sub expires within ~3 days → subscription_expiring notification'),
        ('notify_expired_subscriptions',  'Daily 9:05 AM IST', 'Finds vendors whose sub expired in last 24h → subscription_expired notification'),
    ],
    col_widths=[2.4, 1.6, 2.6],
)

h2('Key Decision')
decision('FCMService moved from chat app to notifications app — originally only chat used push. '
         'After Sprint 11, all 13 apps need push. Single shared FCMService in notifications.fcm '
         'avoids duplication and a circular import.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 12
# ════════════════════════════════════════════════════════════════
sprint_box(12, 'Production Infrastructure + Razorpay Payments')
doc.add_paragraph()

para('Sprint 12 completed two things: (1) full production-ready infrastructure on AWS ECS, '
     'and (2) Razorpay payment integration so vendors can pay for subscriptions through the app.', size=11)

h2('12.1  Production Infrastructure')

h3('Deployment Architecture')
code(
    'Internet (HTTPS)\n'
    '        ↓\n'
    'AWS Application Load Balancer  ← SSL terminated here (AWS ACM certificate)\n'
    '        ↓  HTTP + X-Forwarded-Proto: https\n'
    'Nginx (port 80)  → rate limiting, static served from S3/CDN, WS upgrade\n'
    '        ↓\n'
    'Gunicorn + Uvicorn workers (port 8000)  ← ASGI: handles HTTP + WebSocket\n'
    '        ↓\n'
    'Django (production settings)\n'
    '        ↓\n'
    'PostgreSQL + PostGIS  |  Redis  |  AWS S3  |  Firebase FCM  |  Razorpay'
)

h3('Key Production Settings')
table(
    ['Setting', 'Value', 'Why'],
    [
        ('DEBUG',                      'False',                         'Never expose stack traces in production'),
        ('SECURE_PROXY_SSL_HEADER',    'HTTP_X_FORWARDED_PROTO: https', 'Trust the ALB\'s forwarded HTTPS header'),
        ('SECURE_SSL_REDIRECT',        'False',                         'ALB already redirects HTTP→HTTPS; double-redirect would loop'),
        ('SECURE_HSTS_SECONDS',        '31536000 (1 year)',             'Browser caches HTTPS-only for 1 year after first visit'),
        ('STORAGES.staticfiles',       'S3Boto3Storage',               'Static files served from S3/CDN — not from Django'),
        ('STORAGES.default',           'S3Boto3Storage',               'Media (videos/images) stored in S3'),
        ('CONN_HEALTH_CHECKS',         'True',                         'DB reconnects if connection drops mid-task'),
        ('Logging',                    'JSON to stdout',               'CloudWatch picks up structured JSON automatically'),
        ('Swagger',                    'Blocked at Nginx (403)',        'Never expose API schema on production'),
    ],
    col_widths=[2.0, 1.8, 2.8],
)

h3('scripts/entrypoint.sh')
code(
    '#!/bin/bash\n'
    'set -e\n'
    '# 1. Wait for DB to be ready\n'
    '# 2. python manage.py migrate --noinput\n'
    '# 3. python manage.py collectstatic --noinput --clear\n'
    '# 4. exec gunicorn config.asgi:application --worker-class uvicorn.workers.UvicornWorker\n'
    '\n'
    'This runs on every ECS task start — ensures DB is migrated before serving traffic.\n'
    'Zero-downtime deploys: new task migrates → starts → ALB drains old task.'
)

h3('CI/CD Pipeline (GitHub Actions)')
table(
    ['Stage', 'Trigger', 'Action'],
    [
        ('lint',            'All branch pushes',  'flake8 (style) + black --check (format) + isort --check (imports)'),
        ('test',            'All branch pushes',  'pytest with PostGIS + Redis test containers'),
        ('build',           'main branch only',   'Docker multi-stage build → push to AWS ECR'),
        ('deploy-staging',  'main branch only',   'aws ecs update-service on nearkart-staging cluster (auto)'),
        ('deploy-production','main branch only',  'aws ecs update-service on nearkart-production cluster (manual approval required)'),
    ],
    col_widths=[1.5, 1.4, 3.7],
)

h3('Environments')
table(
    ['Environment', 'Branch', 'Settings Module', 'How to Deploy'],
    [
        ('Local dev',      'any',      'config.settings.development', 'make docker-up'),
        ('CI/CD tests',    'PRs',      'config.settings.testing',     'Automatic on push'),
        ('Staging',        'main',     'config.settings.staging',     'Auto after CI passes'),
        ('Production',     'main',     'config.settings.production',  'Manual approval in GitHub Actions UI'),
    ],
    col_widths=[1.2, 0.8, 2.2, 2.4],
)

divider()
h2('12.2  Razorpay Payment Integration')

para('Before Sprint 12, vendors could only get wallet balance via an admin-initiated topup endpoint. '
     'Sprint 12 added a complete self-service payment flow using Razorpay so vendors can pay '
     'for their own subscription through the app.', size=11)

h3('Payment Flow')
code(
    'Step 1  —  Vendor picks a plan in the app\n'
    '           POST /api/v1/billing/payment/initiate/ {plan_name: "basic"}\n'
    '           ← {order_id, amount: 49900, currency: "INR", razorpay_key_id}\n'
    '\n'
    'Step 2  —  App opens Razorpay checkout SDK with order_id + key_id\n'
    '           User enters card / UPI / net banking → payment processed\n'
    '           ← SDK returns {razorpay_payment_id, razorpay_signature}\n'
    '\n'
    'Step 3  —  App sends payment proof to backend\n'
    '           POST /api/v1/billing/payment/verify/\n'
    '           {razorpay_order_id, razorpay_payment_id, razorpay_signature, plan_name}\n'
    '           → Backend verifies HMAC-SHA256 signature\n'
    '           → Credits wallet with plan price\n'
    '           → Activates subscription\n'
    '           → Sends wallet_topup notification\n'
    '           ← Subscription object returned\n'
    '\n'
    'Backup  —  Razorpay also calls the webhook directly\n'
    '           POST /api/v1/billing/payment/webhook/ (no JWT, verified by signature)\n'
    '           → Idempotency check: if payment_id already in transactions → skip\n'
    '           → Otherwise: same credit + subscribe flow\n'
    '           ← {status: "ok"} or {status: "already_processed"}'
)

h3('Three New Endpoints')
table(
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ('POST', '/billing/payment/initiate/', 'Vendor JWT',     'Create Razorpay order for a plan → get order_id + key_id'),
        ('POST', '/billing/payment/verify/',   'Vendor JWT',     'Verify HMAC → fund wallet → activate subscription'),
        ('POST', '/billing/payment/webhook/',  'None (HMAC sig)','Razorpay webhook backup for payment.captured events'),
    ],
    col_widths=[0.6, 2.2, 1.2, 2.6],
)

h3('Dev Mode (Placeholder Credentials)')
tip('Set RAZORPAY_KEY_ID=rzp_test_PLACEHOLDER in .env (the default).\n'
    '  initiate → returns mock order_id starting with "order_DEV_"\n'
    '  verify → skips HMAC check, any values accepted\n'
    '  webhook → skips signature check\n'
    'No real money moves. Full flow testable without a Razorpay account.')

h3('Going Live')
code(
    '1. Get keys from https://dashboard.razorpay.com/app/keys\n'
    '2. Update .env:\n'
    '     RAZORPAY_KEY_ID=rzp_live_XXXXXXXXXX\n'
    '     RAZORPAY_KEY_SECRET=your_live_secret\n'
    '     RAZORPAY_WEBHOOK_SECRET=your_webhook_secret\n'
    '3. Register webhook URL in Razorpay Dashboard → Settings → Webhooks:\n'
    '     https://api.nearkart.in/api/v1/billing/payment/webhook/\n'
    '     Events: payment.captured\n'
    '4. Test with a small real transaction on staging before production'
)

h3('Key Decisions')
decision('Webhook idempotency via reference_id — verify() stores Razorpay payment_id as Transaction.reference_id. '
         'Webhook checks if reference_id already exists before processing. Safe against Razorpay retry storms.')
decision('Verify credits wallet then subscribes — two separate BillingService calls. '
         'This keeps the same topup+subscribe flow the rest of the system uses. '
         'If subscribe fails, the topup transaction is still recorded for the audit trail.')
decision('No JWT on webhook endpoint (authentication_classes = []) — webhooks come from Razorpay\'s servers, '
         'not from the vendor\'s app. HMAC signature verification replaces JWT.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION — COMPLETE API SURFACE
# ════════════════════════════════════════════════════════════════
h1('Complete API Surface — All 79 Endpoints')

para('Base URL: http://localhost:8000/api/v1   (all endpoints relative)', size=10)

table(
    ['#', 'Method', 'Endpoint', 'Auth', 'Sprint'],
    [
        ('1',  'GET',    '/health/',                                      'None',       'S1'),
        ('2',  'POST',   '/auth/otp/send/',                               'None',       'S2'),
        ('3',  'POST',   '/auth/otp/verify/',                             'None',       'S2'),
        ('4',  'POST',   '/auth/token/refresh/',                          'None',       'S2'),
        ('5',  'GET',    '/auth/me/',                                     'JWT',        'S2'),
        ('6',  'PATCH',  '/auth/me/',                                     'JWT',        'S2'),
        ('7',  'PUT',    '/auth/me/location/',                            'JWT',        'S2'),
        ('8',  'POST',   '/auth/logout/',                                 'JWT',        'S2'),
        ('9',  'GET',    '/auth/users/search/',                           'JWT',        'S10'),
        ('10', 'GET',    '/stores/nearby/',                               'None',       'S3'),
        ('11', 'GET',    '/stores/<uuid>/',                               'None',       'S3'),
        ('12', 'POST',   '/stores/',                                      'Vendor',     'S3'),
        ('13', 'PUT',    '/stores/<uuid>/update/',                        'Owner',      'S3'),
        ('14', 'POST',   '/stores/<uuid>/follow/',                        'JWT',        'S3'),
        ('15', 'POST',   '/stores/<uuid>/review/',                        'JWT',        'S3'),
        ('16', 'GET',    '/stores/<uuid>/qr-code/',                       'Owner',      'S3'),
        ('17', 'GET',    '/stores/<uuid>/hours/',                         'Owner',      'S12'),
        ('18', 'PUT',    '/stores/<uuid>/hours/',                         'Owner',      'S12'),
        ('19', 'POST',   '/stores/<uuid>/blacklist/<uuid>/',              'Vendor',     'S6'),
        ('20', 'GET',    '/stores/<uuid>/blacklist/',                     'Vendor',     'S6'),
        ('21', 'GET',    '/products/nearby/',                             'None',       'S3'),
        ('22', 'GET',    '/products/search/',                             'None',       'S3'),
        ('23', 'GET',    '/products/<uuid>/',                             'None',       'S3'),
        ('24', 'POST',   '/products/',                                    'Vendor',     'S3'),
        ('25', 'PUT',    '/products/<uuid>/update/',                      'Owner',      'S3'),
        ('26', 'DELETE', '/products/<uuid>/update/',                      'Owner',      'S3'),
        ('27', 'POST',   '/products/<uuid>/wishlist/',                    'JWT',        'S3'),
        ('28', 'POST',   '/videos/request-upload/',                       'Vendor',     'S4'),
        ('29', 'POST',   '/videos/<uuid>/confirm-upload/',                'Vendor',     'S4'),
        ('30', 'GET',    '/videos/my-videos/',                            'Vendor',     'S4'),
        ('31', 'GET',    '/videos/feed/',                                 'None',       'S4'),
        ('32', 'GET',    '/videos/<uuid>/',                               'None',       'S4'),
        ('33', 'PATCH',  '/videos/<uuid>/update/',                        'Vendor',     'S4'),
        ('34', 'DELETE', '/videos/<uuid>/delete/',                        'Vendor',     'S4'),
        ('35', 'POST',   '/videos/<uuid>/like/',                          'JWT',        'S4'),
        ('36', 'POST',   '/conversations/start/',                         'JWT',        'S5'),
        ('37', 'GET',    '/conversations/',                               'JWT',        'S5'),
        ('38', 'GET',    '/conversations/<uuid>/messages/',               'JWT',        'S5'),
        ('39', 'PATCH',  '/conversations/<uuid>/read/',                   'JWT',        'S5'),
        ('40', 'GET',    '/billing/plans/',                               'None',       'S7'),
        ('41', 'GET',    '/billing/wallet/',                              'Vendor',     'S7'),
        ('42', 'POST',   '/billing/topup/',                               'Vendor',     'S7'),
        ('43', 'POST',   '/billing/subscribe/',                           'Vendor',     'S7'),
        ('44', 'GET',    '/billing/subscription/',                        'Vendor',     'S7'),
        ('45', 'GET',    '/billing/transactions/',                        'Vendor',     'S7'),
        ('46', 'POST',   '/billing/payment/initiate/',                    'Vendor',     'S12'),
        ('47', 'POST',   '/billing/payment/verify/',                      'Vendor',     'S12'),
        ('48', 'POST',   '/billing/payment/webhook/',                     'None (HMAC)','S12'),
        ('49', 'GET',    '/analytics/vendor/',                            'Vendor',     'S8'),
        ('50', 'GET',    '/analytics/vendor/videos/',                     'Vendor',     'S8'),
        ('51', 'GET',    '/analytics/vendor/products/',                   'Vendor',     'S8'),
        ('52', 'GET',    '/admin-panel/stats/',                           'Staff',      'S8'),
        ('53', 'GET',    '/admin-panel/stores/',                          'Staff',      'S8'),
        ('54', 'PATCH',  '/admin-panel/stores/<uuid>/',                   'Staff',      'S8'),
        ('55', 'GET',    '/admin-panel/users/',                           'Staff',      'S8'),
        ('56', 'POST',   '/admin-panel/users/<uuid>/toggle-active/',      'Staff',      'S8'),
        ('57', 'POST',   '/reservations/',                                'JWT',        'S9'),
        ('58', 'GET',    '/reservations/list/',                           'JWT',        'S9'),
        ('59', 'GET',    '/reservations/<uuid>/',                         'JWT',        'S9'),
        ('60', 'PATCH',  '/reservations/<uuid>/status/',                  'Vendor',     'S9'),
        ('61', 'POST',   '/reservations/<uuid>/cancel/',                  'JWT',        'S9'),
        ('62', 'POST',   '/groups/',                                      'JWT',        'S10'),
        ('63', 'GET',    '/groups/',                                      'JWT',        'S10'),
        ('64', 'GET',    '/groups/<uuid>/',                               'Member',     'S10'),
        ('65', 'DELETE', '/groups/<uuid>/',                               'Creator',    'S10'),
        ('66', 'POST',   '/groups/<uuid>/members/add/',                   'Admin',      'S10'),
        ('67', 'DELETE', '/groups/<uuid>/members/<uuid>/remove/',         'Admin',      'S10'),
        ('68', 'POST',   '/groups/<uuid>/members/<uuid>/make-admin/',     'Admin',      'S10'),
        ('69', 'POST',   '/groups/<uuid>/members/<uuid>/remove-admin/',   'Admin',      'S10'),
        ('70', 'POST',   '/groups/<uuid>/leave/',                         'Member',     'S10'),
        ('71', 'GET',    '/groups/<uuid>/eligible-members/',              'Admin',      'S10'),
        ('72', 'GET',    '/groups/<uuid>/products/',                      'Member',     'S10'),
        ('73', 'POST',   '/groups/<uuid>/products/',                      'Member',     'S10'),
        ('74', 'POST',   '/groups/<uuid>/products/<uuid>/finalize/',      'Admin',      'S10'),
        ('75', 'POST',   '/notifications/device-token/',                  'JWT',        'S11'),
        ('76', 'GET',    '/notifications/',                               'JWT',        'S11'),
        ('77', 'GET',    '/notifications/unread-count/',                  'JWT',        'S11'),
        ('78', 'POST',   '/notifications/<uuid>/read/',                   'JWT',        'S11'),
        ('79', 'POST',   '/notifications/read-all/',                      'JWT',        'S11'),
    ],
    col_widths=[0.3, 0.7, 3.0, 1.0, 0.45],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION — CELERY BEAT SCHEDULE
# ════════════════════════════════════════════════════════════════
h1('Celery Beat — Scheduled Tasks')

para('All scheduled tasks registered in CELERY_BEAT_SCHEDULE in config/settings/base.py.', size=11)

table(
    ['Task Name', 'Schedule', 'What It Does'],
    [
        ('billing.expire_subscriptions',          'Daily midnight IST',   'Marks expired subscriptions as is_active=False'),
        ('reservations.expire_reservations',      'Every hour (:00)',     'Marks pending reservations older than 2h as expired'),
        ('notifications.notify_expiring_subscriptions', 'Daily 9:00 AM IST', 'Sends subscription_expiring notification to vendors expiring in ~3 days'),
        ('notifications.notify_expired_subscriptions',  'Daily 9:05 AM IST', 'Sends subscription_expired notification to vendors who expired in last 24h'),
        ('videos.delete_expired_videos',          'Daily (configurable)', 'Deletes Video records + S3 objects older than VIDEO_EXPIRY_DAYS'),
    ],
    col_widths=[2.6, 1.5, 2.5],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION — DEV MODE BYPASS SYSTEM
# ════════════════════════════════════════════════════════════════
h1('Dev Mode Bypass System')

para('Every external service has a dev-mode bypass. When placeholder or missing credentials are '
     'detected, the system skips the real call and returns mock data. '
     'This means the full app is testable without any paid accounts.', size=11)

table(
    ['Service', 'Dev Mode Trigger', 'Dev Behaviour'],
    [
        ('Twilio (SMS OTP)',  'DEV_FIXED_OTP=123456 in .env',                   'OTP is always 123456, no SMS sent'),
        ('AWS S3 (upload)',   'AWS_ACCESS_KEY_ID contains "EXAMPLE"',            'Presigned URL is mock, video marked ready instantly'),
        ('Firebase FCM',     'FIREBASE_CREDENTIALS_JSON empty or missing',      'Push logged to console: [FCM-DEV] → +91..., not sent'),
        ('Razorpay',         'RAZORPAY_KEY_ID contains "PLACEHOLDER"',          'Mock order_id returned, signatures not verified'),
        ('Sentry',           'SENTRY_DSN empty',                                'SDK not initialised, no errors sent'),
        ('Google Maps',      'GOOGLE_MAPS_API_KEY empty',                       'Locality falls back to "Unknown area"'),
    ],
    col_widths=[1.6, 2.4, 2.6],
)

success('Full end-to-end testing possible in dev:\n'
        '  - OTP login → always 123456\n'
        '  - Video upload → mock URL, instant ready\n'
        '  - Push notifications → logged to console\n'
        '  - Razorpay payment → any values accepted\n'
        'No paid accounts needed for development.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION — QUICK REFERENCE COMMANDS
# ════════════════════════════════════════════════════════════════
h1('Quick Reference Commands')

h2('Start / Stop Dev Stack')
code(
    'make docker-up          # Start all 7 containers\n'
    'make docker-down        # Stop all containers\n'
    'make docker-logs        # Stream logs from all containers\n'
    'make django-shell       # Open Django shell inside container\n'
    'make migrate            # Run migrations\n'
    'make seed-plans         # Seed Free/Basic/Premium plans\n'
    'make create-superuser   # Create Django admin superuser'
)

h2('Start / Stop Production Stack')
code(
    'make prod-up            # Start production stack (docker-compose.prod.yml)\n'
    'make prod-down          # Stop production stack\n'
    'make prod-logs          # Stream production logs (JSON format)\n'
    'make check-env          # Verify all required env vars are set'
)

h2('Test the Stack')
code(
    '# Health check\n'
    'curl http://localhost:8000/api/v1/health/\n'
    '\n'
    '# Send OTP\n'
    'curl -X POST http://localhost:8000/api/v1/auth/otp/send/ \\\n'
    '  -H "Content-Type: application/json" \\\n'
    '  -d \'{"phone_number": "+919999999999"}\'\n'
    '\n'
    '# Verify OTP (dev OTP is always 123456)\n'
    'curl -X POST http://localhost:8000/api/v1/auth/otp/verify/ \\\n'
    '  -H "Content-Type: application/json" \\\n'
    '  -d \'{"phone_number": "+919999999999", "otp": "123456"}\'\n'
    '\n'
    '# Swagger UI (dev only)\n'
    'open http://localhost:8000/api/docs/\n'
    '\n'
    '# Django Admin\n'
    'open http://localhost:8000/admin/'
)

h2('Regenerate This Document')
code(
    'python3 docs/gen_project_journey.py\n'
    '# Overwrites: docs/NearKart_Complete_Project_Journey_Sprint0_to_Sprint12.docx'
)

# ── FOOTER ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'NearKart Backend — Complete Project Journey   |   Sprint 0 to Sprint 12   |   May 2026   |   79 Endpoints Built'
)
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── SAVE ──
out = '/Users/hazeevali/Documents/NearKart/Backend/nearkart_backend/docs/NearKart_Complete_Project_Journey_Sprint0_to_Sprint12.docx'
doc.save(out)
print(f'Saved: {out}')

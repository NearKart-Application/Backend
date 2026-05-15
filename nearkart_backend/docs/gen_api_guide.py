from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── PAGE MARGINS ──
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(0.9)
section.top_margin  = section.bottom_margin = Inches(0.9)

# ── HELPERS ──
def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def shade_para(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def h1(text):
    return doc.add_heading(text, level=1)

def h2(text):
    return doc.add_heading(text, level=2)

def h3(text):
    return doc.add_heading(text, level=3)

def para(text, bold=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def code(text, indent=0.2):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    shade_para(p, 'F0F0F0')
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x8C)
    return p

def note(text, color='FFF3CD', text_color=(0x7D, 0x4E, 0x00)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    shade_para(p, color)
    run = p.add_run('  ' + text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(*text_color)
    return p

def success(text):
    return note('✅  ' + text, 'D4EDDA', (0x15, 0x57, 0x24))

def warning(text):
    return note('⚠️  ' + text, 'FFF3CD', (0x7D, 0x4E, 0x00))

def tip(text):
    return note('💡  ' + text, 'D1ECF1', (0x0C, 0x54, 0x60))

def error_box(text):
    return note('🚫  ' + text, 'F8D7DA', (0x7B, 0x1D, 0x1D))

def table(headers, rows, col_widths=None, header_color='1F4E79', alt_color='DEEAF1'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, header_color)
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            if ri % 2 == 0:
                shade_cell(cell, alt_color)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def method_badge(method):
    colors = {
        'GET':    ('D4EDDA', (0x15, 0x55, 0x24)),
        'POST':   ('CCE5FF', (0x00, 0x4D, 0xAA)),
        'PUT':    ('FFF3CD', (0x7D, 0x4E, 0x00)),
        'PATCH':  ('FFF3CD', (0x7D, 0x4E, 0x00)),
        'DELETE': ('F8D7DA', (0x7B, 0x1D, 0x1D)),
    }
    return colors.get(method, ('EEEEEE', (0x33, 0x33, 0x33)))

def endpoint_header(method, url, auth, sprint):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    bg, fg = method_badge(method)
    r1 = p.add_run(f'  {method}  ')
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(*fg)
    r2 = p.add_run(f'  {url}')
    r2.font.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    r3 = p.add_run(f'   [{auth}]   [{sprint}]')
    r3.font.size = Pt(9)
    r3.font.italic = True
    r3.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    shade_para(p, bg)
    return p

def divider():
    p = doc.add_paragraph('─' * 90)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0]
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

# ════════════════════════════════════════════════════════════════
#  COVER
# ════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('NearKart Backend')
run.font.name = 'Calibri'; run.font.size = Pt(30); run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Master API Testing Guide')
run2.font.size = Pt(18); run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('All endpoints · Request bodies · Auth headers · Step-by-step flows')
run3.font.size = Pt(11); run3.font.italic = True; run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run('Sprint 1 → Sprint 12   |   Last updated: May 2026')
run4.font.size = Pt(10); run4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
note('📌  MAINTAINER NOTE: Every time a new sprint adds APIs, update this document.\n'
     '     Add the new section under the correct sprint heading and update the Quick Reference table.',
     'D1ECF1', (0x0C, 0x54, 0x60))

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 1 — SETUP
# ════════════════════════════════════════════════════════════════
h1('1. One-Time Setup')

h2('1.1  Start the Server')
para('Run these commands every time you want to test:')
code('cd /Users/hazeevali/Documents/NearKart/Backend/nearkart_backend\ndocker compose up -d')
para('Verify server is running:')
code('curl http://localhost:8000/api/v1/health/')
success('Expected: {"status": "ok", "db": "ok", "redis": "ok"}')
warning('If db or redis shows "error" → run: docker compose up -d')

h2('1.2  Base URL')
code('http://localhost:8000/api/v1')
para('All endpoint URLs in this guide are relative to this base URL.', size=10)

h2('1.3  Swagger UI (Browser)')
code('http://localhost:8000/api/docs/')
tip('Use Swagger to READ what endpoints exist and what fields they need.\n'
    'Use Postman or curl for actual testing — Swagger has display bugs with expired tokens.\n\n'
    'How to use Swagger:\n'
    '  1. Open http://localhost:8000/api/docs/ in browser\n'
    '  2. Each endpoint has pre-filled example request bodies — "Try it out" is ON by default\n'
    '     (no need to click "Try it out" manually — just click Execute)\n'
    '  3. For protected endpoints: first run OTP send + verify → copy "access" token\n'
    '  4. Click the lock icon "Authorize" (top right of page)\n'
    '     → Paste JUST the token in the Value field — do NOT add "Bearer " prefix\n'
    '       Swagger adds "Bearer " automatically for you\n'
    '     → Click Authorize → Close\n'
    '  5. All protected endpoints will now work — token survives page refresh\n\n'
    'Swagger improvements active:\n'
    '  • tryItOutEnabled=true  — "Try it out" pre-clicked on all endpoints\n'
    '  • persistAuthorization=true  — your token survives page refresh\n'
    '  • defaultModelsExpandDepth=-1  — schema models panel collapsed (less clutter)\n'
    '  • Pre-filled examples on: OTP send/verify, Store create/review, Product create/update,\n'
    '    Video upload/confirm/update, Start Conversation')

error_box('SWAGGER UUID PLACEHOLDER — VERY COMMON MISTAKE:\n\n'
          '     Swagger auto-fills a dummy UUID in every <uuid> path field:\n'
          '       3fa85f64-5717-4562-b3fc-2c963f66afa6\n\n'
          '     This UUID does NOT exist in your database.\n'
          '     If you click Execute without replacing it, you will always get:\n'
          '       404 — "not_found" or "Conversation not found" or "Video not found"\n\n'
          '     CORRECT WORKFLOW in Swagger:\n'
          '       Step 1 — Run the CREATE endpoint first  (e.g. POST /conversations/start/)\n'
          '       Step 2 — Copy the "id" field from that response\n'
          '       Step 3 — Paste it into the UUID field of the next endpoint\n'
          '       Step 4 — Then click Execute\n\n'
          '     Examples of what to replace:\n'
          '       /conversations/{conversation_id}/messages/  → paste real conversation id\n'
          '       /videos/{video_id}/                         → paste real video id\n'
          '       /stores/{id}/                               → paste real store id\n'
          '       /products/{id}/                             → paste real product id')

# ── PHONE NUMBER FORMAT RULES ──
h2('1.4  Phone Number Format Rules  (IMPORTANT)')
warning('Every phone number in NearKart must follow the exact Indian mobile format.\n'
        'A single wrong character causes a 400 validation error.')

para('Format pattern:', bold=True, size=10)
code('+91XXXXXXXXXX\n\nWhere:\n  +91    = India country code (required, with the + sign)\n  X      = 10 digits\n  First digit must be 6, 7, 8, or 9 (Indian mobile range)')

para('Valid examples:', bold=True, size=10)
table(
    ['Phone Number', 'Why It Is Valid'],
    [
        ('+919999999999', 'Starts with +91, 10 digits, first digit is 9'),
        ('+916000000001', 'Starts with +91, 10 digits, first digit is 6'),
        ('+917500000000', 'Starts with +91, 10 digits, first digit is 7'),
        ('+918123456789', 'Starts with +91, 10 digits, first digit is 8'),
    ],
    col_widths=[2.0, 4.5],
    header_color='155724',
    alt_color='D4EDDA',
)

para('Invalid examples — these will return 400 error:', bold=True, size=10)
table(
    ['Wrong Input', 'Error Reason', 'Correct Version'],
    [
        ('9999999999',      'Missing +91 prefix',                    '+919999999999'),
        ('919999999999',    'Missing + sign (has 91 but no +)',       '+919999999999'),
        ('+91 9999999999',  'Space between +91 and digits',           '+919999999999'),
        ('+91-9999999999',  'Hyphen between +91 and digits',          '+919999999999'),
        ('+915000000000',   'First digit is 5 — must be 6, 7, 8, 9', '+916000000000'),
        ('+911234567890',   'First digit is 1 — must be 6, 7, 8, 9', '+919234567890'),
        ('+9199999999',     'Only 8 digits after +91 — need 10',      '+919999999999'),
        ('+919999999999 ',  'Trailing space at end',                   '+919999999999'),
    ],
    col_widths=[1.8, 2.2, 1.8],
    header_color='721C24',
    alt_color='F8D7DA',
)

error_box('Copy-paste phone numbers carefully. A space, missing +, or wrong first digit = 400 error.\n'
          '     The API returns: {"error": "validation_error", "message": "Enter a valid Indian mobile number in +91XXXXXXXXXX format."}')

h2('1.5  OTP Rules (Dev Mode)')
table(
    ['Rule', 'Value / Behaviour'],
    [
        ('Dev OTP (always)',         '123456  — hardcoded in .env as DEV_FIXED_OTP'),
        ('OTP expiry',               '10 minutes after calling /otp/send/'),
        ('Max wrong attempts',       '5 attempts — after 5 wrong OTPs the session is locked'),
        ('Rate limit on /otp/send/', '5 requests per 5 minutes per IP'),
        ('New user on first login',  'Yes — user is auto-created with role based on phone number'),
        ('OTP in production',        'Real SMS via Twilio (Sprint 12) — not active yet'),
    ],
    col_widths=[2.2, 4.3],
)
tip('In Dev mode: always use OTP 123456. No SMS is sent, no Twilio needed.')
warning('Call /otp/send/ BEFORE /otp/verify/ — if no OTP was sent, verify returns "Invalid OTP".')

h2('1.6  Token Rules')
table(
    ['Token Type', 'Lifetime', 'Where to Get It', 'How to Use It'],
    [
        ('access',  '1 hour',   'POST /auth/otp/verify/ → "access" field',  'Authorization: Bearer <access_token>'),
        ('refresh', '30 days',  'POST /auth/otp/verify/ → "refresh" field', 'Body of POST /auth/token/refresh/'),
    ],
    col_widths=[0.9, 0.8, 2.6, 2.2],
)
warning('When access token expires (after 1 hour): call POST /auth/token/refresh/ with your refresh token.\n'
        'You get a new access token without re-entering OTP.\n'
        'When refresh token expires (after 30 days): must do OTP send + verify again.')
error_box('401 token_invalid  = your access token expired → call /auth/token/refresh/\n'
          '     401 authentication_failed  = Authorization header missing or malformed → add "Bearer " before token')

h2('1.7  Setup Postman Environment (Do Once)')
para('1.  Open Postman → Environments → + → Name: NearKart Local')
para('2.  Add these variables:')
table(
    ['Variable', 'Initial Value', 'How It Gets Set'],
    [
        ('base_url',        'http://localhost:8000/api/v1', 'Set manually once'),
        ('vendor_token',    '(empty)',                       'Auto-filled by OTP verify request'),
        ('vendor_refresh',  '(empty)',                       'Auto-filled by OTP verify request'),
        ('customer_token',  '(empty)',                       'Auto-filled by OTP verify request'),
        ('refresh_token',   '(empty)',                       'Auto-filled by OTP verify request'),
        ('store_id',          '(empty)',                       'Copy from store create response'),
        ('product_id',        '(empty)',                       'Copy from product create response'),
        ('video_id',          '(empty)',                       'Auto-filled by request-upload request'),
        ('conversation_id',     '(empty)',                       'Auto-filled by start conversation request'),
        ('notification_id',     '(empty)',                       'Auto-filled by list notifications request'),
        ('razorpay_order_id',   '(empty)',                       'Auto-filled by payment initiate request'),
        ('reservation_id',      '(empty)',                       'Auto-filled by create reservation request'),
        ('group_id',            '(empty)',                       'Auto-filled by create group request'),
        ('sp_id',               '(empty)',                       'Auto-filled by share product request'),
    ],
    col_widths=[1.5, 2.2, 2.8]
)
para('3.  Select NearKart Local from the environment dropdown (top-right in Postman)')
para('4.  In the OTP Verify request → Tests tab → paste this script:')
code('// Auto-save tokens after login\nconst r = pm.response.json();\nif (r.access) {\n    pm.environment.set("vendor_token", r.access);   // change to customer_token for customer\n    pm.environment.set("vendor_refresh", r.refresh);\n    console.log("Token saved:", r.access.substring(0, 30) + "...");\n}')

h2('1.8  Auth Header Format')
warning('Token expires in 1 hour. If you get 401, run OTP send + verify again to get a fresh token.')
table(
    ['Situation', 'Authorization Header Value'],
    [
        ('Public endpoint (no auth needed)', '(leave empty — do not add Authorization header)'),
        ('Any logged-in user',               'Bearer {{vendor_token}}  or  Bearer {{customer_token}}'),
        ('Vendor-only endpoint',             'Bearer {{vendor_token}}'),
        ('Store owner endpoint',             'Bearer {{vendor_token}}  (must be the owner of that store)'),
    ],
    col_widths=[2.5, 4.0]
)
error_box('Common mistakes:\n'
          '     ✗  "bearer eyJhbGci..."  (lowercase b)   — must be capital B: Bearer\n'
          '     ✗  "eyJhbGci..."  (no "Bearer " prefix)  — must have "Bearer " before the token\n'
          '     ✗  "Bearer  eyJhbGci..."  (two spaces)    — must be exactly one space\n'
          '     ✓  "Bearer eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9..."  (correct)')

h2('1.9  Common Error Codes Quick Reference')
table(
    ['HTTP Code', 'Error Key', 'Cause', 'Fix'],
    [
        ('200', 'OK',                      '—', 'Success'),
        ('201', 'Created',                 '—', 'Resource created successfully'),
        ('204', 'No Content',              '—', 'Deleted successfully (empty body)'),
        ('400', 'validation_error',        'Missing or invalid field in request body', 'Check the "details" key in response — it names the exact field that failed'),
        ('400', 'otp_invalid',             'Wrong OTP, expired OTP, or OTP not sent yet', 'Call /otp/send/ first, then verify with OTP 123456 in dev'),
        ('400', 'validation_error (phone)','Phone not in +91XXXXXXXXXX format', 'See Section 1.4 — Phone Number Format Rules'),
        ('401', 'authentication_failed',   'No Authorization header, or "Bearer " prefix missing', 'Add header: Authorization: Bearer <your_token>'),
        ('401', 'token_invalid',           'Access token expired or refresh token blacklisted', 'Call /auth/token/refresh/ for new access token, or re-login'),
        ('403', 'permission_denied',       'Wrong role (e.g. customer on vendor endpoint) or not store owner', 'Use correct role token'),
        ('404', 'not_found',               'UUID does not exist, resource is inactive, wrong URL, OR Swagger placeholder UUID used', 'Check UUID is real (not 3fa85f64-...). Copy id from the CREATE response first.'),
        ('429', 'throttled',               'Too many OTP requests (5 per 5 min)', 'Wait before sending another OTP'),
    ],
    col_widths=[0.8, 1.6, 2.2, 2.0]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 2 — COMPLETE API REFERENCE
# ════════════════════════════════════════════════════════════════
h1('2. Complete API Reference')

table(
    ['#', 'Method', 'Endpoint', 'Auth Required', 'Sprint'],
    [
        ('1',  'GET',    '/health/',                    'None (public)',   'S1'),
        ('2',  'POST',   '/auth/otp/send/',             'None (public)',   'S2'),
        ('3',  'POST',   '/auth/otp/verify/',           'None (public)',   'S2'),
        ('4',  'POST',   '/auth/token/refresh/',        'None (public)',   'S2'),
        ('5',  'GET',    '/auth/me/',                   'Bearer JWT',      'S2'),
        ('6',  'PATCH',  '/auth/me/',                   'Bearer JWT',      'S2'),
        ('7',  'PUT',    '/auth/me/location/',          'Bearer JWT',      'S2'),
        ('8',  'POST',   '/auth/logout/',               'Bearer JWT',      'S2'),
        ('9',  'GET',    '/stores/nearby/',             'None (public)',   'S3'),
        ('10', 'GET',    '/stores/<uuid>/',             'None (public)',   'S3'),
        ('11', 'POST',   '/stores/',                    'Vendor JWT only', 'S3'),
        ('12', 'PUT',    '/stores/<uuid>/update/',      'Store Owner JWT', 'S3'),
        ('13', 'POST',   '/stores/<uuid>/follow/',      'Bearer JWT',      'S3'),
        ('14', 'POST',   '/stores/<uuid>/review/',      'Bearer JWT',      'S3'),
        ('15', 'GET',    '/stores/<uuid>/qr-code/',     'Store Owner JWT', 'S3'),
        ('16', 'GET',    '/products/nearby/',           'None (public)',   'S3'),
        ('17', 'GET',    '/products/search/',           'None (public)',   'S3'),
        ('18', 'GET',    '/products/<uuid>/',           'None (public)',   'S3'),
        ('19', 'POST',   '/products/',                  'Vendor JWT only', 'S3'),
        ('20', 'PUT',    '/products/<uuid>/update/',    'Store Owner JWT', 'S3'),
        ('21', 'DELETE', '/products/<uuid>/update/',    'Store Owner JWT', 'S3'),
        ('22', 'POST',   '/products/<uuid>/wishlist/',  'Bearer JWT',      'S3'),
        ('23', 'POST',   '/videos/request-upload/',          'Vendor JWT only', 'S4'),
        ('24', 'POST',   '/videos/<uuid>/confirm-upload/',   'Vendor JWT only', 'S4'),
        ('25', 'GET',    '/videos/my-videos/',               'Vendor JWT only', 'S4'),
        ('26', 'PATCH',  '/videos/<uuid>/update/',           'Vendor JWT only', 'S4'),
        ('27', 'GET',    '/videos/feed/',                    'None (public)',   'S4'),
        ('28', 'GET',    '/videos/<uuid>/',                  'None (public)',   'S4'),
        ('29', 'DELETE', '/videos/<uuid>/delete/',              'Vendor JWT only', 'S4'),
        ('30', 'POST',   '/videos/<uuid>/like/',              'Bearer JWT',      'S4'),
        ('31', 'POST',   '/conversations/start/',                          'Bearer JWT',      'S5'),
        ('32', 'GET',    '/conversations/',                               'Bearer JWT',      'S5'),
        ('33', 'GET',    '/conversations/<uuid>/messages/',               'Bearer JWT',      'S5'),
        ('34', 'PATCH',  '/conversations/<uuid>/read/',                   'Bearer JWT',      'S5'),
        ('35', 'POST',   '/stores/<uuid>/blacklist/<uuid>/',              'Vendor JWT only', 'S6'),
        ('36', 'GET',    '/stores/<uuid>/blacklist/',                     'Vendor JWT only', 'S6'),
        ('37', 'GET',    '/billing/plans/',                               'None (public)',   'S7'),
        ('38', 'GET',    '/billing/wallet/',                              'Vendor JWT only', 'S7'),
        ('39', 'POST',   '/billing/topup/',                               'Vendor JWT only', 'S7'),
        ('40', 'POST',   '/billing/subscribe/',                           'Vendor JWT only', 'S7'),
        ('41', 'GET',    '/billing/subscription/',                        'Vendor JWT only', 'S7'),
        ('42', 'GET',    '/billing/transactions/',                        'Vendor JWT only', 'S7'),
        ('43', 'GET',    '/analytics/vendor/',                            'Vendor JWT only', 'S8'),
        ('44', 'GET',    '/analytics/vendor/videos/',                    'Vendor JWT only', 'S8'),
        ('45', 'GET',    '/analytics/vendor/products/',                  'Vendor JWT only', 'S8'),
        ('46', 'GET',    '/admin-panel/stats/',                          'Staff JWT only',  'S8'),
        ('47', 'GET',    '/admin-panel/stores/',                         'Staff JWT only',  'S8'),
        ('48', 'PATCH',  '/admin-panel/stores/<uuid>/',                  'Staff JWT only',  'S8'),
        ('49', 'GET',    '/admin-panel/users/',                          'Staff JWT only',  'S8'),
        ('50', 'POST',   '/admin-panel/users/<uuid>/toggle-active/',     'Staff JWT only',  'S8'),
        ('51', 'POST',   '/reservations/',                               'Bearer JWT',      'S9'),
        ('52', 'GET',    '/reservations/list/',                          'Bearer JWT',      'S9'),
        ('53', 'GET',    '/reservations/<uuid>/',                        'Bearer JWT',      'S9'),
        ('54', 'PATCH',  '/reservations/<uuid>/status/',                 'Vendor JWT only', 'S9'),
        ('55', 'POST',   '/reservations/<uuid>/cancel/',                 'Bearer JWT',      'S9'),
        ('56', 'GET',    '/auth/users/search/',                          'Bearer JWT',      'S10'),
        ('57', 'POST',   '/groups/',                                     'Bearer JWT',      'S10'),
        ('58', 'GET',    '/groups/',                                     'Bearer JWT',      'S10'),
        ('59', 'GET',    '/groups/<uuid>/',                              'Member JWT',      'S10'),
        ('60', 'DELETE', '/groups/<uuid>/',                              'Creator JWT',     'S10'),
        ('61', 'POST',   '/groups/<uuid>/members/add/',                  'Admin JWT',       'S10'),
        ('62', 'DELETE', '/groups/<uuid>/members/<uuid>/remove/',        'Admin JWT',       'S10'),
        ('63', 'POST',   '/groups/<uuid>/members/<uuid>/make-admin/',    'Admin JWT',       'S10'),
        ('64', 'POST',   '/groups/<uuid>/members/<uuid>/remove-admin/',  'Admin JWT',       'S10'),
        ('65', 'POST',   '/groups/<uuid>/leave/',                        'Member JWT',      'S10'),
        ('66', 'GET',    '/groups/<uuid>/eligible-members/',             'Admin JWT',       'S10'),
        ('67', 'GET',    '/groups/<uuid>/products/',                     'Member JWT',      'S10'),
        ('68', 'POST',   '/groups/<uuid>/products/',                     'Member JWT',      'S10'),
        ('69', 'POST',   '/groups/<uuid>/products/<uuid>/finalize/',     'Admin JWT',       'S10'),
        ('70', 'POST',   '/notifications/device-token/',                 'Bearer JWT',      'S11'),
        ('71', 'GET',    '/notifications/',                              'Bearer JWT',      'S11'),
        ('72', 'GET',    '/notifications/unread-count/',                 'Bearer JWT',      'S11'),
        ('73', 'POST',   '/notifications/<uuid>/read/',                  'Bearer JWT',      'S11'),
        ('74', 'POST',   '/notifications/read-all/',                     'Bearer JWT',      'S11'),
        ('75', 'GET',    '/stores/<uuid>/hours/',                        'Store Owner JWT', 'S12'),
        ('76', 'PUT',    '/stores/<uuid>/hours/',                        'Store Owner JWT', 'S12'),
        ('77', 'POST',   '/billing/payment/initiate/',                   'Vendor JWT only', 'S12'),
        ('78', 'POST',   '/billing/payment/verify/',                     'Vendor JWT only', 'S12'),
        ('79', 'POST',   '/billing/payment/webhook/',                    'None (signature)','S12'),
    ],
    col_widths=[0.3, 0.7, 2.8, 1.4, 0.4]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 3 — STEP BY STEP TEST FLOW
# ════════════════════════════════════════════════════════════════
h1('3. Step-by-Step Test Flow (Full End-to-End)')
para('Follow these steps in order. Each step builds on the previous one.', size=11)
note('📌  Dev OTP is always 123456 — no real SMS needed.\n'
     '     Phone must be +91 followed by 10 digits, first digit 6–9. See Section 1.4 for full rules.')

# ── STEP 1 ──
divider()
h2('STEP 1 — Health Check')
endpoint_header('GET', '{{base_url}}/health/', 'No Auth', 'Sprint 1')
para('No body needed. No auth needed.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "status": "ok",\n  "db": "ok",\n  "redis": "ok",\n  "version": "1.0.0",\n  "environment": "development"\n}')
success('If this works → server is running, DB and Redis are connected.')
warning('If db or redis shows "error" → run: docker compose up -d')

# ── STEP 2 ──
divider()
h2('STEP 2 — Send OTP (Vendor)')
endpoint_header('POST', '{{base_url}}/auth/otp/send/', 'No Auth', 'Sprint 2')
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body:', bold=True, size=10)
code('{\n  "phone_number": "+919999999999"\n}')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "message": "OTP sent successfully"\n}')
tip('Phone number rules (see full table in Section 1.4):\n'
    '  ✓  +919999999999  (correct — +91, first digit 9)\n'
    '  ✓  +916000000001  (correct — +91, first digit 6)\n'
    '  ✗  9999999999     (wrong — missing +91)\n'
    '  ✗  919999999999   (wrong — missing + sign)\n'
    '  ✗  +91 9999999999 (wrong — space not allowed)\n'
    '  ✗  +915000000000  (wrong — first digit must be 6, 7, 8, or 9)')
error_box('400 Error if format wrong:\n'
          '     {"error": "validation_error",\n'
          '      "message": "Enter a valid Indian mobile number in +91XXXXXXXXXX format.",\n'
          '      "code": "ERROR",\n'
          '      "details": {"phone_number": ["Enter a valid Indian mobile number in +91XXXXXXXXXX format."]}}')

# ── STEP 3 ──
divider()
h2('STEP 3 — Verify OTP → Get Vendor Token')
endpoint_header('POST', '{{base_url}}/auth/otp/verify/', 'No Auth', 'Sprint 2')
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body:', bold=True, size=10)
code('{\n  "phone_number": "+919999999999",\n  "otp": "123456"\n}')
tip('In dev mode OTP is always 123456.\n'
    'Must call /otp/send/ first — if no OTP was sent, verify returns "Invalid OTP".\n'
    'After 5 wrong OTP attempts the session is locked — call /otp/send/ again.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "message": "Login successful",\n  "access": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...",\n  "refresh": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...",\n  "user": {\n    "id": "5f9a17a1-b9c9-436a-bebb-65c9057ea803",\n    "phone_number": "+919999999999",\n    "role": "vendor",\n    "full_name": "",\n    "email": "",\n    "created_at": "2026-05-15T10:02:00Z"\n  }\n}')
success('Copy the "access" value → this is your vendor_token.\nCopy the "refresh" value → save as vendor_refresh (use for token refresh later).')
tip('In Postman Tests tab paste:\n'
    '  const r = pm.response.json();\n'
    '  pm.environment.set("vendor_token", r.access);\n'
    '  pm.environment.set("vendor_refresh", r.refresh);\n\n'
    'In Swagger UI: copy the "access" value → click Authorize (lock icon, top right)\n'
    '  → paste JUST the token — do NOT add "Bearer " prefix (Swagger adds it automatically)\n'
    '  → Click Authorize → Close — protected endpoints now work for 1 hour.')
warning('Role is determined by phone number used on FIRST login:\n'
        '  +919999999999 = vendor  |  +916000000001 = customer\n'
        'Roles are assigned in Django admin — check user.role if unexpected.')

# ── STEP 4 ──
divider()
h2('STEP 4 — Get Customer Token')
endpoint_header('POST', '{{base_url}}/auth/otp/send/', 'No Auth', 'Sprint 2')
para('Request Body:', bold=True, size=10)
code('{\n  "phone_number": "+916000000001"\n}')
para('Then verify:')
endpoint_header('POST', '{{base_url}}/auth/otp/verify/', 'No Auth', 'Sprint 2')
code('{\n  "phone_number": "+916000000001",\n  "otp": "123456"\n}')
success('Copy "access" value → save as customer_token.')
warning('Use a DIFFERENT phone number from the vendor.\n'
        'Each phone number belongs to one user with one role.\n'
        'A vendor cannot act as a customer with the same phone.')

# ── STEP 5 ──
divider()
h2('STEP 5 — Get My Profile')
endpoint_header('GET', '{{base_url}}/auth/me/', 'Bearer JWT', 'Sprint 2')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('No request body needed.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "id": "5f9a17a1-b9c9-436a-bebb-65c9057ea803",\n  "phone_number": "+919999999999",\n  "role": "vendor",\n  "full_name": "",\n  "email": "",\n  "created_at": "2026-05-15T10:02:00Z"\n}')
error_box('No Authorization header → 401 authentication_failed\n'
          '     Expired token → 401 token_invalid → call /auth/token/refresh/ or re-login')

# ── STEP 6 ──
divider()
h2('STEP 6 — Update Profile')
endpoint_header('PATCH', '{{base_url}}/auth/me/', 'Bearer JWT', 'Sprint 2')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body (send only fields you want to change):', bold=True, size=10)
code('{\n  "full_name": "Rahul Kumar",\n  "email": "rahul@example.com"\n}')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "id": "...",\n  "phone_number": "+919999999999",\n  "role": "vendor",\n  "full_name": "Rahul Kumar",\n  "email": "rahul@example.com",\n  "created_at": "..."\n}')
tip('Only full_name and email can be updated.\n'
    'phone_number and role are read-only — they cannot be changed via this endpoint.')

# ── STEP 7 ──
divider()
h2('STEP 7 — Update Location')
endpoint_header('PUT', '{{base_url}}/auth/me/location/', 'Bearer JWT', 'Sprint 2')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body:', bold=True, size=10)
code('{\n  "latitude": 13.0827,\n  "longitude": 80.2707\n}')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "message": "Location updated"\n}')
tip('latitude range: -90.0 to 90.0   |   longitude range: -180.0 to 180.0\n'
    'Chennai example: latitude 13.0827, longitude 80.2707')
error_box('latitude > 90 or < -90  → 400 validation_error\n'
          '     longitude > 180 or < -180  → 400 validation_error')

# ── STEP 8 ──
divider()
h2('STEP 8 — Refresh Access Token')
endpoint_header('POST', '{{base_url}}/auth/token/refresh/', 'No Auth', 'Sprint 2')
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body:', bold=True, size=10)
code('{\n  "refresh": "{{vendor_refresh}}"\n}')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "access": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9..."\n}')
tip('When to use: access token expires after 1 hour.\n'
    'Use this endpoint to get a new access token without re-entering OTP.\n'
    'Get your refresh token from the /otp/verify/ response → "refresh" field.\n'
    'Refresh token is valid for 30 days.')
error_box('Refresh token already used after logout → 401 token_invalid\n'
          '     Refresh token expired (30 days) → 401 token_invalid → re-do full OTP login')

# ── STEP 9 ──
divider()
h2('STEP 9 — Logout')
endpoint_header('POST', '{{base_url}}/auth/logout/', 'Bearer JWT', 'Sprint 2')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body:', bold=True, size=10)
code('{\n  "refresh": "{{vendor_refresh}}"\n}')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "message": "Logged out successfully"\n}')
success('After logout: the refresh token is blacklisted permanently.\n'
        'Using it again in /token/refresh/ returns 401 token_invalid.\n'
        'The access token still works until it naturally expires (up to 1 hour).')

doc.add_page_break()

# ── STEP 10 ──
h2('STEP 10 — Create Store  (Vendor only)')
endpoint_header('POST', '{{base_url}}/stores/', 'Vendor JWT', 'Sprint 3')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body:', bold=True, size=10)
code('{\n  "name": "Fashion Hub",\n  "description": "Trendy clothes for everyone",\n  "category": "fashion",\n  "phone": "+919876543210",\n  "address": "123 Anna Salai, Chennai",\n  "latitude": 13.0418,\n  "longitude": 80.2341,\n  "logo_url": "https://example.com/logo.png",\n  "banner_url": "https://example.com/banner.png"\n}')
tip('Swagger has a pre-filled sample body for this endpoint ("Sample store (Chennai)").\n'
    'Open Swagger → POST /stores/ → the body is pre-filled with name, address, coords — just click Execute.')
para('Required fields: name, address, latitude, longitude', size=10, indent=0.2)
para('Optional fields: description, category, phone, logo_url, banner_url', size=10, indent=0.2)
para('Valid category values:', size=10, indent=0.2)
code('fashion | jewellery | footwear | decor | furniture | gifts | beauty | food | electronics | other')
para('Phone field in store body: same +91XXXXXXXXXX format rule applies (see Section 1.4)', size=10, indent=0.2)
para('Expected Response  201 Created:', bold=True, size=10)
code('{\n  "id": "6c8adfdd-a788-4661-88e7-0768a037745e",\n  "owner_phone": "+919999999999",\n  "name": "Fashion Hub",\n  "category": "fashion",\n  "locality": "Anna Salai",\n  "lat": 13.0418,\n  "lng": 80.2341,\n  "is_verified": false,\n  "is_open": false,\n  "performance_score": 0.0,\n  "follower_count": 0,\n  "hours": [],\n  "created_at": "..."\n}')
success('Copy "id" from response → save as store_id.')
warning('is_verified is false by default → store will NOT appear in nearby results yet.\n'
        'To enable for testing:\n'
        '  Option A: Django Admin → http://localhost:8000/admin/ → Stores → tick is_verified → Save\n'
        '  Option B: Run in terminal:\n'
        '    docker compose exec django /venv/bin/python manage.py shell -c \\\n'
        '    "from apps.stores.models import Store; Store.objects.filter(name=\'Fashion Hub\').update(is_verified=True)"')

# ── STEP 11 ──
divider()
h2('STEP 11 — Nearby Stores  (Public)')
endpoint_header('GET', '{{base_url}}/stores/nearby/', 'No Auth', 'Sprint 3')
para('No body. Query parameters:', bold=True, size=10)
table(
    ['Parameter', 'Required', 'Default', 'Description'],
    [
        ('lat',      'Yes', '—',   'User latitude (e.g. 13.0418 for Chennai)'),
        ('lng',      'Yes', '—',   'User longitude (e.g. 80.2341 for Chennai)'),
        ('radius',   'No',  '2',   'Search radius in km. Allowed values: 1, 2, 3, 5'),
        ('category', 'No',  'all', 'Filter by category slug (e.g. fashion)'),
    ],
    col_widths=[1.2, 0.8, 0.8, 3.7]
)
para('Full URL examples:', size=10)
code('GET {{base_url}}/stores/nearby/?lat=13.0418&lng=80.2341&radius=2\nGET {{base_url}}/stores/nearby/?lat=13.0418&lng=80.2341&radius=2&category=fashion')
para('Expected Response  200 OK:', bold=True, size=10)
code('[\n  {\n    "id": "6c8adfdd-a788-4661-88e7-0768a037745e",\n    "name": "Fashion Hub",\n    "category": "fashion",\n    "locality": "Anna Salai",\n    "logo_url": "",\n    "is_open": false,\n    "is_verified": true,\n    "performance_score": 0.0,\n    "lat": 13.0418,\n    "lng": 80.2341,\n    "distance_km": 0.12\n  }\n]')
warning('Returns [] if no verified stores found in radius.\n'
        'Make sure store is_verified=true (see Step 10 note).\n'
        'Missing lat or lng → 400 validation_error')

# ── STEP 12 ──
divider()
h2('STEP 12 — Store Detail  (Public)')
endpoint_header('GET', '{{base_url}}/stores/{{store_id}}/', 'No Auth', 'Sprint 3')
para('No body. Replace {{store_id}} with the UUID from Step 10.')
warning('In Swagger: replace the placeholder UUID 3fa85f64-... with the real store_id from Step 10.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "id": "6c8adfdd-...",\n  "name": "Fashion Hub",\n  "description": "...",\n  "category": "fashion",\n  "phone": "+919876543210",\n  "address": "123 Anna Salai",\n  "locality": "Anna Salai",\n  "lat": 13.0418,\n  "lng": 80.2341,\n  "logo_url": "",\n  "banner_url": "",\n  "is_verified": true,\n  "is_open": false,\n  "performance_score": 0.0,\n  "follower_count": 0,\n  "hours": [],\n  "distance_km": null,\n  "created_at": "..."\n}')
tip('Response is cached in Redis for 5 minutes.\n'
    'After updating a store, wait up to 5 minutes for the detail endpoint to reflect changes.')

# ── STEP 13 ──
divider()
h2('STEP 13 — Update Store  (Owner Only)')
endpoint_header('PUT', '{{base_url}}/stores/{{store_id}}/update/', 'Owner JWT', 'Sprint 3')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body (partial — send only fields you want to change):', bold=True, size=10)
code('{\n  "is_open": true,\n  "description": "New description here"\n}')
para('Any store field can be updated:', size=10, indent=0.2)
code('name | description | category | phone | address | latitude | longitude |\nlogo_url | banner_url | is_open | is_active')
para('Expected Response  200 OK:', bold=True, size=10)
code('{ ... full updated store object ... }')
error_box('Using a different vendor\'s token → 403 permission_denied\n'
          '     Using a customer token → 403 permission_denied')

# ── STEP 14 ──
divider()
h2('STEP 14 — Follow / Unfollow Store  (Toggle)')
endpoint_header('POST', '{{base_url}}/stores/{{store_id}}/follow/', 'Bearer JWT', 'Sprint 3')
para('Authorization: Bearer {{customer_token}}', size=10, indent=0.2)
para('No request body needed — do not send any JSON.')
para('Expected Response  200 OK (first call — follows):', bold=True, size=10)
code('{\n  "followed": true,\n  "message": "Following store."\n}')
para('Expected Response  200 OK (second call — unfollows):', bold=True, size=10)
code('{\n  "followed": false,\n  "message": "Unfollowed store."\n}')
success('This is a toggle. First call follows, second call unfollows. No separate unfollow endpoint.')

# ── STEP 15 ──
divider()
h2('STEP 15 — Add / Update Review  (Any Logged-in User)')
endpoint_header('POST', '{{base_url}}/stores/{{store_id}}/review/', 'Bearer JWT', 'Sprint 3')
para('Authorization: Bearer {{customer_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body:', bold=True, size=10)
code('{\n  "rating": 5,\n  "comment": "Excellent store! Great products."\n}')
tip('Swagger has two pre-filled examples for this endpoint: "Five-star review" and "Three-star review".\n'
    'Click the dropdown in Swagger to switch between them.')
para('rating: required integer 1 to 5', size=10, indent=0.2)
para('comment: optional string', size=10, indent=0.2)
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "id": "4117a307-...",\n  "user_phone": "+916000000001",\n  "rating": 5,\n  "comment": "Excellent store!",\n  "created_at": "..."\n}')
success('Calling again with the same user + store UPDATES the existing review — no duplicate created.\n'
        'Store\'s performance_score auto-recalculates after every review.')
error_box('rating 0 or 6 → 400 validation_error — rating must be 1 to 5')

# ── STEP 16 ──
divider()
h2('STEP 16 — QR Code  (Store Owner Only)')
endpoint_header('GET', '{{base_url}}/stores/{{store_id}}/qr-code/', 'Owner JWT', 'Sprint 3')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('No request body needed.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "qr_code_url": "https://cdn.nearkart.in/qrcodes/<store_id>/qr.png"\n}')
tip('In development: qr_code_url will be empty string "" because AWS S3 is not configured.\n'
    'This works fully in production with real AWS credentials (Sprint 12).')

doc.add_page_break()

# ── STEP 17 ──
h2('STEP 17 — Create Product  (Vendor Only)')
endpoint_header('POST', '{{base_url}}/products/', 'Vendor JWT', 'Sprint 3')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body:', bold=True, size=10)
code('{\n  "name": "Cotton Kurta",\n  "description": "Handwoven cotton kurta",\n  "category": "fashion",\n  "status": "active",\n  "is_visible": true,\n  "base_price": "499.00",\n  "variants": [\n    {\n      "name": "Small",\n      "sku": "KT-S-001",\n      "price": "499.00",\n      "stock_quantity": 10\n    },\n    {\n      "name": "Medium",\n      "sku": "KT-M-001",\n      "price": "499.00",\n      "stock_quantity": 8\n    },\n    {\n      "name": "Large",\n      "sku": "KT-L-001",\n      "price": "549.00",\n      "stock_quantity": 5\n    }\n  ]\n}')
tip('Swagger has two pre-filled examples for this endpoint:\n'
    '  "Kurta with variants" — fashion product with 3 size variants\n'
    '  "Simple product (no variants)" — food product with no variants\n'
    'Select from the Examples dropdown in Swagger before clicking Execute.\n'
    'Note: SKUs must be globally unique — if you run this more than once, change the SKU values.')
para('Required fields: name, base_price', size=10, indent=0.2)
para('status options: draft | active | inactive | out_of_stock', size=10, indent=0.2)
para('variants: optional array. Each variant needs: name, sku (globally unique), price, stock_quantity', size=10, indent=0.2)
para('Expected Response  201 Created:', bold=True, size=10)
code('{\n  "id": "f7e32e6b-...",\n  "store_id": "6c8adfdd-...",\n  "store_name": "Fashion Hub",\n  "name": "Cotton Kurta",\n  "status": "active",\n  "base_price": "499.00",\n  "variants": [ ... ],\n  "images": [],\n  "is_wishlisted": false,\n  "created_at": "...",\n  "last_updated_at": "..."\n}')
success('Copy "id" → save as product_id.')
warning('Vendor must have a store first (Step 10). If not: 400 "Create a store first".\n'
        'Product with status=draft will NOT appear in nearby/search results.\n'
        'SKU must be globally unique — duplicate SKU → 400 validation_error.')

# ── STEP 18 ──
divider()
h2('STEP 18 — Nearby Products  (Public)')
endpoint_header('GET', '{{base_url}}/products/nearby/', 'No Auth', 'Sprint 3')
para('No body. Query parameters:', bold=True, size=10)
table(
    ['Parameter', 'Required', 'Default', 'Description'],
    [
        ('lat',      'Yes', '—',   'User latitude'),
        ('lng',      'Yes', '—',   'User longitude'),
        ('radius',   'No',  '2',   'Radius in km. Allowed: 1, 2, 3, 5'),
        ('category', 'No',  'all', 'Filter by product category'),
    ],
    col_widths=[1.2, 0.8, 0.8, 3.7]
)
code('GET {{base_url}}/products/nearby/?lat=13.0418&lng=80.2341&radius=2')
para('Expected Response  200 OK:', bold=True, size=10)
code('[\n  {\n    "id": "f7e32e6b-...",\n    "store_name": "Fashion Hub",\n    "name": "Cotton Kurta",\n    "category": "fashion",\n    "base_price": "499.00",\n    "min_price": "499.00",\n    "primary_image": null,\n    "distance_km": 0.12,\n    "status": "active"\n  }\n]')
warning('ALL of these must be true for a product to appear:\n'
        '  • Product status = active\n'
        '  • Product is_visible = true\n'
        '  • Store is_verified = true\n'
        '  • At least one variant with stock_quantity > 0')

# ── STEP 19 ──
divider()
h2('STEP 19 — Search Products  (Public)')
endpoint_header('GET', '{{base_url}}/products/search/', 'No Auth', 'Sprint 3')
para('No body. Query parameters:', bold=True, size=10)
table(
    ['Parameter', 'Required', 'Description'],
    [
        ('q',      'Yes', 'Search term — partial name matches work (min 20% trigram similarity)'),
        ('lat',    'No',  'Latitude — filters results to this location when combined with lng'),
        ('lng',    'No',  'Longitude — filters results to this location when combined with lat'),
        ('radius', 'No',  'Search radius in km when lat/lng provided (default: 5)'),
    ],
    col_widths=[1.0, 0.8, 4.7]
)
code('GET {{base_url}}/products/search/?q=kurta\nGET {{base_url}}/products/search/?q=kurta&lat=13.0418&lng=80.2341&radius=5')
success('Results sorted by similarity — closest match first.\n'
        'Partial matches work: "kurt" matches "Cotton Kurta". Returns max 30 results.')
error_box('Missing ?q= → 400 validation_error: "Search query q is required."')

# ── STEP 20 ──
divider()
h2('STEP 20 — Product Detail  (Public)')
endpoint_header('GET', '{{base_url}}/products/{{product_id}}/', 'No Auth', 'Sprint 3')
para('No body. Replace {{product_id}} with UUID from Step 17.')
warning('In Swagger: replace the placeholder UUID 3fa85f64-... with the real product_id from Step 17.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "id": "f7e32e6b-...",\n  "store_id": "...",\n  "store_name": "Fashion Hub",\n  "name": "Cotton Kurta",\n  "status": "active",\n  "base_price": "499.00",\n  "variants": [ ... ],\n  "images": [],\n  "is_wishlisted": false,\n  "created_at": "...",\n  "last_updated_at": "..."\n}')
tip('Add Authorization: Bearer {{customer_token}} header → is_wishlisted field reflects actual wishlist state.')
warning('Only products with status=active AND is_visible=true are returned.\n'
        'All others (draft, inactive, hidden) → 404 not_found.')

# ── STEP 21 ──
divider()
h2('STEP 21 — Update Product  (Owner Only)')
endpoint_header('PUT', '{{base_url}}/products/{{product_id}}/update/', 'Owner JWT', 'Sprint 3')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body (partial — send only fields to change):', bold=True, size=10)
code('{\n  "base_price": "449.00",\n  "status": "active"\n}')
tip('Swagger has two pre-filled examples for this endpoint:\n'
    '  "Update price and visibility" — changes base_price and is_visible\n'
    '  "Mark product inactive" — sets status=inactive and is_visible=false')
para('Any product field can be updated:', size=10, indent=0.2)
code('name | description | category | status | is_visible | base_price')
para('Expected Response  200 OK — full updated product object.')
error_box('Note: variants cannot be updated via this endpoint after creation.\n'
          '     To change variants: delete and re-create the product (Sprint 4 will add variant-level update).')

# ── STEP 22 ──
divider()
h2('STEP 22 — Delete Product  (Owner Only)')
endpoint_header('DELETE', '{{base_url}}/products/{{product_id}}/update/', 'Owner JWT', 'Sprint 3')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('No request body needed.')
para('Expected Response  204 No Content  (empty body — no JSON returned)')
success('Product is permanently deleted. GET on the same UUID returns 404 after this.')

# ── STEP 23 ──
divider()
h2('STEP 23 — Wishlist Toggle  (Any Logged-in User)')
endpoint_header('POST', '{{base_url}}/products/{{product_id}}/wishlist/', 'Bearer JWT', 'Sprint 3')
para('Authorization: Bearer {{customer_token}}', size=10, indent=0.2)
para('No request body needed — do not send any JSON.')
para('Expected Response  200 OK (first call — adds):', bold=True, size=10)
code('{ "wishlisted": true, "message": "Added to wishlist." }')
para('Expected Response  200 OK (second call — removes):', bold=True, size=10)
code('{ "wishlisted": false, "message": "Removed from wishlist." }')
success('Toggle — first call adds, second call removes. No separate remove endpoint.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 4 — VIDEO MODULE
# ════════════════════════════════════════════════════════════════
h1('Sprint 4 — Video Module')
para('Follow Steps 24–31 after completing Step 23. Vendor must have a store (Step 10).', size=11)
note('📌  Dev mode: No real AWS S3 needed. Mock upload URLs are returned. '
     'Celery transcoding is instant — video goes ready in ~2 seconds.')

# ── STEP 24 ──
divider()
h2('STEP 24 — Request Presigned Upload URL  (Vendor only)')
endpoint_header('POST', '{{base_url}}/videos/request-upload/', 'Vendor JWT', 'Sprint 4')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body:', bold=True, size=10)
code('{\n  "title": "Summer Kurta Collection",\n  "description": "Handwoven cotton kurtas for the season"\n}')
para('Required: title (max 200 chars)', size=10, indent=0.2)
para('Optional: description', size=10, indent=0.2)
para('Expected Response  201 Created:', bold=True, size=10)
code('{\n  "video_id": "82d5ede8-d051-4ad9-9054-c922a87a3773",\n  "upload_url": "https://mock-s3.dev/videos/raw/.../original.mp4?dev=true",\n  "expires_in_seconds": 900,\n  "message": "Upload URL ready. PUT your video file to upload_url..."\n}')
success('Copy "video_id" → save as video_id.\n'
        'In production: PUT the video file to upload_url using Content-Type: video/mp4.\n'
        'In dev: skip the PUT — just call confirm-upload directly.')
tip('The presigned URL expires in 900 seconds (15 minutes).\n'
    'If it expires, call request-upload again to get a new URL.')
error_box('No store created yet → 400: "Create a store first before uploading videos."\n'
          '     Customer token → 403: Vendor access only.\n'
          '     Missing title → 400: This field is required.')

# ── STEP 25 ──
divider()
h2('STEP 25 — Confirm Upload → Trigger Transcoding  (Vendor only)')
endpoint_header('POST', '{{base_url}}/videos/{{video_id}}/confirm-upload/', 'Vendor JWT', 'Sprint 4')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body:', bold=True, size=10)
code('{\n  "duration_seconds": 45\n}')
para('duration_seconds: optional (0–60). Set to the actual video length.', size=10, indent=0.2)
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "video_id": "82d5ede8-...",\n  "status": "processing",\n  "message": "Transcoding queued. Check GET /videos/<id>/ for status updates."\n}')
tip('Dev mode: Celery immediately marks the video ready with mock URLs.\n'
    'Wait 2–3 seconds, then GET /videos/<id>/ to verify status = ready.')
error_box('Call confirm-upload again on a ready video → 400: "Video is already in ready state."\n'
          '     Using another vendor\'s video_id → 404 not_found.')

# ── STEP 26 ──
divider()
h2('STEP 26 — My Videos  (Vendor Library)')
endpoint_header('GET', '{{base_url}}/videos/my-videos/', 'Vendor JWT', 'Sprint 4')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('No request body. Optional query parameter:', bold=True, size=10)
table(
    ['Parameter', 'Required', 'Values', 'Description'],
    [
        ('status', 'No', 'pending_upload | processing | ready | failed | expired', 'Filter to a specific status'),
    ],
    col_widths=[1.0, 0.8, 2.6, 2.1]
)
code('GET {{base_url}}/videos/my-videos/\nGET {{base_url}}/videos/my-videos/?status=ready')
para('Expected Response  200 OK:', bold=True, size=10)
code('[\n  {\n    "id": "82d5ede8-...",\n    "title": "Summer Kurta Collection",\n    "status": "ready",\n    "video_url": "https://mock-s3.dev/...",\n    "duration_seconds": 45,\n    "view_count": 1,\n    "like_count": 0,\n    "is_visible": true,\n    "expires_at": "2026-06-14T...",\n    "created_at": "2026-05-15T..."\n  }\n]')
tip('Use ?status=processing after confirm-upload to monitor transcoding progress.\n'
    'Returns [] (empty array) if vendor has no store yet — no error.')
error_box('Customer token → 403 Vendor access only.')

# ── STEP 27 ──
divider()
h2('STEP 27 — Update Video  (Vendor Only)')
endpoint_header('PATCH', '{{base_url}}/videos/{{video_id}}/update/', 'Vendor JWT', 'Sprint 4')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body (send only the fields you want to change):', bold=True, size=10)
code('{\n  "title": "Revised Summer Collection",\n  "is_visible": false\n}')
para('Patchable fields: title, description, is_visible', size=10, indent=0.2)
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "id": "82d5ede8-...",\n  "title": "Revised Summer Collection",\n  "is_visible": false,\n  ...\n}')
tip('Set is_visible: false to hide a video from the public feed without deleting it.\n'
    'Set is_visible: true to re-publish it.')
error_box('Another vendor\'s video_id → 404 not_found (filtered to own store).\n'
          '     Customer token → 403 Vendor access only.')

# ── STEP 28 ──
divider()
h2('STEP 28 — Video Detail  (Public)')
endpoint_header('GET', '{{base_url}}/videos/{{video_id}}/', 'No Auth', 'Sprint 4')
para('No body. Replace {{video_id}} with UUID from Step 24.')
warning('In Swagger: replace the placeholder UUID 3fa85f64-... with the real video_id from Step 24.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "id": "82d5ede8-...",\n  "store_id": "6c8adfdd-...",\n  "store_name": "Fashion Hub",\n  "title": "Summer Kurta Collection",\n  "description": "Handwoven cotton kurtas for the season",\n  "video_url": "https://mock-s3.dev/videos/hls/.../master.m3u8?dev=true",\n  "thumbnail_url": "https://mock-s3.dev/videos/thumbnails/.../thumb.jpg?dev=true",\n  "status": "ready",\n  "duration_seconds": 45,\n  "view_count": 1,\n  "like_count": 0,\n  "is_liked": false,\n  "locality": "Anna Salai",\n  "distance_km": null,\n  "is_visible": true,\n  "expires_at": "2026-06-14T...",\n  "created_at": "2026-05-15T..."\n}')
tip('view_count increments by 1 on every GET call.\n'
    'Add Authorization header to see is_liked = true/false for the current user.')
warning('Only videos with status=ready AND is_visible=true are returned.\n'
        'Videos still processing → 404 not_found (check back after transcoding).')

# ── STEP 29 ──
divider()
h2('STEP 29 — Video Feed  (Public — Location-based)')
endpoint_header('GET', '{{base_url}}/videos/feed/', 'No Auth', 'Sprint 4')
para('No body. Query parameters:', bold=True, size=10)
table(
    ['Parameter', 'Required', 'Default', 'Description'],
    [
        ('lat',      'Yes', '—',  'User latitude (e.g. 13.0418 for Chennai)'),
        ('lng',      'Yes', '—',  'User longitude (e.g. 80.2341 for Chennai)'),
        ('radius',   'No',  '5',  'Search radius in km (default 5, no max limit)'),
        ('store_id', 'No',  'all','Filter to a specific store UUID'),
    ],
    col_widths=[1.2, 0.8, 0.6, 3.9]
)
code('GET {{base_url}}/videos/feed/?lat=13.0418&lng=80.2341&radius=10\nGET {{base_url}}/videos/feed/?lat=13.0418&lng=80.2341&store_id={{store_id}}')
para('Expected Response  200 OK:', bold=True, size=10)
code('[\n  {\n    "id": "82d5ede8-...",\n    "store_name": "Fashion Hub",\n    "title": "Summer Kurta Collection",\n    "video_url": "https://mock-s3.dev/...",\n    "thumbnail_url": "https://mock-s3.dev/...",\n    "status": "ready",\n    "duration_seconds": 45,\n    "view_count": 2,\n    "like_count": 0,\n    "distance_km": 0.45\n  }\n]')
success('Results sorted by distance first, then newest. Max 50 videos per call.')
warning('Only ready + visible + non-expired videos appear.\n'
        'Missing lat or lng → 400 validation_error.')

# ── STEP 30 ──
divider()
h2('STEP 30 — Like / Unlike Video  (Toggle)')
endpoint_header('POST', '{{base_url}}/videos/{{video_id}}/like/', 'Bearer JWT', 'Sprint 4')
para('Authorization: Bearer {{customer_token}}', size=10, indent=0.2)
para('No request body needed — do not send any JSON.')
para('Expected Response  200 OK (first call — like):', bold=True, size=10)
code('{\n  "liked": true,\n  "message": "Liked."\n}')
para('Expected Response  200 OK (second call — unlike):', bold=True, size=10)
code('{\n  "liked": false,\n  "message": "Unliked."\n}')
success('Toggle — first call likes, second call unlikes. No separate unlike endpoint.')
tip('The video\'s like_count updates atomically using SQL F() expressions — safe under concurrent requests.')

# ── STEP 31 ──
divider()
h2('STEP 31 — Delete Video  (Store Owner Only)')
endpoint_header('DELETE', '{{base_url}}/videos/{{video_id}}/delete/', 'Vendor JWT', 'Sprint 4')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('No request body needed.')
para('Expected Response  204 No Content  (empty body)')
success('Video permanently deleted. GET on the same UUID returns 404 after this.')
error_box('Customer token → 403 Vendor access only.\n'
          '     Different vendor\'s video_id → 404 (filtered to own store, not 403).')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 5 — CHAT MODULE
# ════════════════════════════════════════════════════════════════
h1('Sprint 5 — Chat Module')
para('Follow Steps 32–36 after completing Step 23 (store created). Both vendor and customer tokens needed.', size=11)
note('📌  WebSocket port is 8001 (Daphne), not 8000 (Django REST).\n'
     '     Token is passed as a query param: ?token=<jwt_access_token>\n'
     '     Dev mode: FCM push notifications are logged, not actually sent.')

# ── STEP 32 ──
divider()
h2('STEP 32 — Start / Get Conversation')
endpoint_header('POST', '{{base_url}}/conversations/start/', 'Bearer JWT', 'Sprint 5')
para('Authorization: Bearer {{customer_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body:', bold=True, size=10)
code('{\n  "store_id": "{{store_id}}"\n}')
para('Expected Response  201 Created (first time) / 200 OK (already exists):', bold=True, size=10)
code('{\n  "id": "68fe68f7-d330-46cc-94e7-7853caab6d54",\n  "store_id": "6c8adfdd-...",\n  "store_name": "Fashion Hub",\n  "customer_phone": "+916000000001",\n  "my_unread_count": 0,\n  "last_message": null,\n  "last_message_at": null,\n  "is_active": true,\n  "created_at": "2026-05-15T..."\n}')
success('Copy "id" → save as conversation_id.\n'
        '201 = new conversation created. 200 = existing conversation returned (idempotent).')
tip('Vendor can also call this endpoint with customer_id to open a chat:\n'
    '  {"customer_id": "<customer_uuid>"}')
error_box('Missing store_id → 400: "store_id is required."\n'
          '     Invalid store → 404: Store not found.\n'
          '     No auth header → 401: authentication_failed.')

# ── STEP 33 ──
divider()
h2('STEP 33 — Connect WebSocket + Send Message  (Real-time)')
para('WebSocket URL:', bold=True, size=10)
code('ws://localhost:8001/ws/conversations/{{conversation_id}}/?token={{customer_token}}')
warning('Port is 8001 (Daphne/ASGI), not 8000.\n'
        'Token goes in the query string — NOT in a header.')
para('After connecting, send this JSON message:', bold=True, size=10)
code('{"type": "chat_message", "content": "Hello from customer!"}')
para('Expected broadcast received by ALL connected participants:', bold=True, size=10)
code('{\n  "id": "448b97e3-...",\n  "conversation_id": "68fe68f7-...",\n  "sender_id": "fbe9358d-...",\n  "sender_phone": "+916000000001",\n  "sender_role": "customer",\n  "content": "Hello from customer!",\n  "message_type": "text",\n  "media_url": "",\n  "ref_id": null,\n  "is_read": false,\n  "created_at": "2026-05-15T..."\n}')
success('Message is saved to DB and broadcast in real-time.\n'
        'If vendor is offline → FCM push sent to their registered device tokens.\n'
        'Dev mode: push is logged, not sent (no Firebase credentials needed).')
tip('Test with Python:\n\n'
    '  import asyncio, json, websockets\n\n'
    '  async def chat():\n'
    '      uri = "ws://localhost:8001/ws/conversations/<id>/?token=<jwt>"\n'
    '      async with websockets.connect(uri) as ws:\n'
    '          await ws.send(json.dumps({"type": "chat_message", "content": "Hello!"}))\n'
    '          print(json.loads(await ws.recv()))\n\n'
    '  asyncio.run(chat())\n\n'
    'Or use Postman v10.10+ WebSocket tab.\n'
    'Or use CLI:  wscat -c "ws://localhost:8001/ws/conversations/<id>/?token=<jwt>"')
error_box('No token → WS closes with code 4001 (unauthorized).\n'
          '     Not a member of the conversation → WS closes with code 4003 (forbidden).\n'
          '     Empty content → ignored silently. Unknown type → ignored silently.')

# ── STEP 34 ──
divider()
h2('STEP 34 — List Conversations (Inbox)')
endpoint_header('GET', '{{base_url}}/conversations/', 'Bearer JWT', 'Sprint 5')
para('Authorization: Bearer {{customer_token}}  OR  Bearer {{vendor_token}}', size=10, indent=0.2)
para('No body. No query params.')
para('Expected Response  200 OK:', bold=True, size=10)
code('[\n  {\n    "id": "68fe68f7-...",\n    "store_id": "6c8adfdd-...",\n    "store_name": "Fashion Hub",\n    "customer_phone": "+916000000001",\n    "my_unread_count": 1,\n    "last_message": {\n      "id": "448b97e3-...",\n      "content": "Hello from customer!",\n      "sender_role": "customer",\n      "created_at": "2026-05-15T..."\n    },\n    "last_message_at": "2026-05-15T...",\n    "is_active": true,\n    "created_at": "2026-05-15T..."\n  }\n]')
tip('Customer sees all their conversations across all stores.\n'
    'Vendor sees all conversations for their store (entire inbox).\n'
    'Sorted newest first by last_message_at.')

# ── STEP 35 ──
divider()
h2('STEP 35 — Message History (Paginated)')
endpoint_header('GET', '{{base_url}}/conversations/{{conversation_id}}/messages/', 'Bearer JWT', 'Sprint 5')
warning('In Swagger: replace the placeholder UUID 3fa85f64-... with the real conversation_id from Step 32.')
para('Authorization: Bearer {{customer_token}}', size=10, indent=0.2)
para('No body. Optional query parameter:', bold=True, size=10)
table(
    ['Parameter', 'Required', 'Description'],
    [
        ('before', 'No', 'Message UUID — returns up to 50 messages older than this one (for pagination)'),
    ],
    col_widths=[1.0, 0.8, 4.7]
)
code('GET {{base_url}}/conversations/{{conversation_id}}/messages/\nGET {{base_url}}/conversations/{{conversation_id}}/messages/?before={{oldest_message_id}}')
para('Expected Response  200 OK:', bold=True, size=10)
code('[\n  {\n    "id": "448b97e3-...",\n    "conversation_id": "68fe68f7-...",\n    "sender_id": "fbe9358d-...",\n    "sender_phone": "+916000000001",\n    "sender_role": "customer",\n    "content": "Hello from customer!",\n    "message_type": "text",\n    "media_url": "",\n    "ref_id": null,\n    "is_read": false,\n    "created_at": "2026-05-15T..."\n  }\n]')
tip('Returns max 50 messages, oldest first.\n'
    'For infinite scroll: take the id of the first (oldest) message in the list → pass as ?before=<id>.')
error_box('Accessing another user\'s conversation → 403: permission_denied.\n'
          '     Non-existent conversation_id → 404: not_found.')

# ── STEP 36 ──
divider()
h2('STEP 36 — Mark Conversation as Read')
endpoint_header('PATCH', '{{base_url}}/conversations/{{conversation_id}}/read/', 'Bearer JWT', 'Sprint 5')
warning('In Swagger: replace the placeholder UUID 3fa85f64-... with the real conversation_id from Step 32.')
para('Authorization: Bearer {{vendor_token}}  (or customer token)', size=10, indent=0.2)
para('No request body needed — do not send any JSON.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "message": "Marked as read."\n}')
success('Resets the caller\'s unread_count to 0.\n'
        'Marks all received messages (from the other party) as is_read=true in DB.')
tip('Call this when the user opens the chat screen in the app.\n'
    'After this call, GET /conversations/ will show my_unread_count = 0 for this conversation.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 6 — BLACKLIST ENGINE
# ════════════════════════════════════════════════════════════════
h1('Sprint 6 — Blacklist Engine')
para('Follow Steps 37–38 after completing Step 10 (Create Store). Vendor must own a store.', size=11)
tip('Blacklist is per store — a vendor blocks a specific customer from interacting with their store.\n'
    'A blocked customer cannot: follow the store, review it, start a new chat, or connect via WebSocket.\n'
    'Blocking and unblocking both use the same endpoint — it toggles on each call.')

# ── STEP 37 ──
divider()
h2('STEP 37 — Block / Unblock a Customer  (Toggle)')
endpoint_header('POST', '{{base_url}}/stores/{{store_id}}/blacklist/{{customer_id}}/', 'Vendor JWT', 'Sprint 6')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body (reason is optional):', bold=True, size=10)
code('{\n  "reason": "Spamming and abusive messages"\n}')
tip('Swagger has two pre-filled examples: "Block with reason" and "Block no reason".\n'
    'Replace {{store_id}} with real store UUID and {{customer_id}} with the customer UUID\n'
    '(get customer UUID from GET /auth/me/ using the customer token).')
para('Expected Response  200 OK — first call (blocks):', bold=True, size=10)
code('{\n  "is_blocked": true,\n  "message": "Customer blocked."\n}')
para('Expected Response  200 OK — second call (unblocks):', bold=True, size=10)
code('{\n  "is_blocked": false,\n  "message": "Customer unblocked."\n}')
success('Toggle — first call blocks, second call unblocks. No separate unblock endpoint.')
warning('Only the vendor who OWNS the store can block customers.\n'
        'Trying to block from a store you don\'t own → 403 permission_denied.\n'
        'The customer_id must belong to a user with role=customer → 404 if not found.')
error_box('403 blacklisted — what a blocked customer gets when they try to:\n'
          '     POST /stores/<id>/follow/       → 403 — You cannot follow this store.\n'
          '     POST /stores/<id>/review/       → 403 — You cannot review this store.\n'
          '     POST /conversations/start/      → 403 — You cannot start a conversation with this store.\n'
          '     WebSocket connect               → close code 4003')

# ── STEP 38 ──
divider()
h2('STEP 38 — List Blocked Customers')
endpoint_header('GET', '{{base_url}}/stores/{{store_id}}/blacklist/', 'Vendor JWT', 'Sprint 6')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('No request body needed.')
para('Expected Response  200 OK:', bold=True, size=10)
code('[\n  {\n    "id": "92b1f816-a34a-488b-9c36-867e020db838",\n    "customer_phone": "+916000000001",\n    "customer_name": "",\n    "reason": "Spamming and abusive messages",\n    "created_at": "2026-05-15T17:31:25Z"\n  }\n]')
success('Returns [] when no customers are blocked.')
warning('In Swagger: replace the placeholder UUID 3fa85f64-... with the real store_id from Step 10.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 7 — BILLING + WALLET
# ════════════════════════════════════════════════════════════════
h1('Sprint 7 — Billing + Wallet')
para('Follow Steps 39–44 after completing Step 10 (Create Store). All billing endpoints are vendor-only except Plans.', size=11)

table(
    ['Plan', 'Price', 'Videos', 'Products'],
    [
        ('Free',    '₹0 / month',   '3',         '10'),
        ('Basic',   '₹499 / month', '20',        '50'),
        ('Premium', '₹999 / month', 'Unlimited', 'Unlimited'),
    ],
    col_widths=[1.2, 1.4, 1.4, 1.4],
    header_color='1F4E79',
)
tip('Top up your wallet first (Step 41), then subscribe to a plan (Step 42).\n'
    'The Free plan costs ₹0 — you can subscribe without any balance.\n'
    'Plan limits are enforced when uploading videos (Step 24) and creating products (Step 17).')

# ── STEP 39 ──
divider()
h2('STEP 39 — List Plans  (Public)')
endpoint_header('GET', '{{base_url}}/billing/plans/', 'No Auth', 'Sprint 7')
para('No body. No auth required.')
para('Expected Response  200 OK:', bold=True, size=10)
code('[\n  {\n    "name": "free",\n    "display_name": "Free Plan",\n    "price": "0.00",\n    "duration_days": 30,\n    "video_limit": 3,\n    "product_limit": 10,\n    "video_limit_display": "3",\n    "product_limit_display": "10",\n    "description": "Up to 3 videos and 10 products. No payment required."\n  },\n  { "name": "basic", "price": "499.00", "video_limit_display": "20", ... },\n  { "name": "premium", "price": "999.00", "video_limit_display": "Unlimited", ... }\n]')

# ── STEP 40 ──
divider()
h2('STEP 40 — Check Wallet Balance  (Vendor Only)')
endpoint_header('GET', '{{base_url}}/billing/wallet/', 'Vendor JWT', 'Sprint 7')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('No body.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "store_name": "Fashion Hub",\n  "wallet_balance": "0.00"\n}')

# ── STEP 41 ──
divider()
h2('STEP 41 — Top Up Wallet  (Vendor Only)')
endpoint_header('POST', '{{base_url}}/billing/topup/', 'Vendor JWT', 'Sprint 7')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body:', bold=True, size=10)
code('{\n  "amount": "1000.00"\n}')
tip('Swagger has two pre-filled examples: "Top up ₹500" and "Top up ₹1000".\n'
    'Dev mode: money is added instantly — no payment gateway.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "message": "₹1000.00 added to wallet.",\n  "amount_added": "1000.00",\n  "wallet_balance": "1000.00",\n  "transaction_id": "55a8c23d-..."\n}')
error_box('400 — amount must be a positive number  (if amount is 0 or negative)\n'
          '     400 — validation_error  (if amount is not a number, e.g. "abc")')

# ── STEP 42 ──
divider()
h2('STEP 42 — Subscribe to a Plan  (Vendor Only)')
endpoint_header('POST', '{{base_url}}/billing/subscribe/', 'Vendor JWT', 'Sprint 7')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body:', bold=True, size=10)
code('{\n  "plan_name": "basic"\n}')
tip('Swagger has three pre-filled examples: "Subscribe Basic", "Subscribe Premium", "Subscribe Free".\n'
    'Valid plan_name values: free | basic | premium\n'
    'Subscribing again (same or different plan) replaces the current subscription.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "id": "e3f0d941-...",\n  "store_name": "Fashion Hub",\n  "plan": {\n    "name": "basic",\n    "display_name": "Basic Plan",\n    "price": "499.00",\n    "video_limit": 20,\n    "product_limit": 50\n  },\n  "started_at": "2026-05-15T19:24:45Z",\n  "expires_at": "2026-06-14T19:24:45Z",\n  "is_active": true,\n  "days_left": 29\n}')
success('Wallet balance is deducted automatically. GET /billing/wallet/ will show updated balance.')
error_box('400 — insufficient_balance  (not enough in wallet)\n'
          '     Fix: top up wallet first (Step 41), then retry subscribe.\n'
          '     404 — Plan not found  (wrong plan_name)\n'
          '     Fix: use exactly one of: free | basic | premium')

# ── STEP 43 ──
divider()
h2('STEP 43 — Subscription Status  (Vendor Only)')
endpoint_header('GET', '{{base_url}}/billing/subscription/', 'Vendor JWT', 'Sprint 7')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('No body.')
para('Expected Response  200 OK — same shape as subscribe response.', size=10)
tip('days_left shows how many days remain.\n'
    'If is_active=false, the plan has expired — subscribe again to reactivate.\n'
    'A vendor with no subscription gets 404 from this endpoint.')

# ── STEP 44 ──
divider()
h2('STEP 44 — Transaction History  (Vendor Only)')
endpoint_header('GET', '{{base_url}}/billing/transactions/', 'Vendor JWT', 'Sprint 7')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('No body.')
para('Expected Response  200 OK:', bold=True, size=10)
code('[\n  {\n    "id": "...",\n    "type": "subscription",\n    "amount": "-499.00",\n    "description": "Subscribed to Basic Plan",\n    "reference_id": "SUB-BASIC-1747315485",\n    "balance_after": "501.00",\n    "created_at": "2026-05-15T19:24:45Z"\n  },\n  {\n    "type": "topup",\n    "amount": "1000.00",\n    "description": "Wallet top-up of ₹1000.00",\n    "reference_id": "DEV-TOPUP-1747315200",\n    "balance_after": "1000.00"\n  }\n]')
tip('Top-up entries have positive amount. Subscription entries have negative amount.\n'
    'Sorted newest first. Free plan subscriptions create NO transaction (no money moved).')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 8 — ANALYTICS + ADMIN PANEL
# ════════════════════════════════════════════════════════════════
h1('Sprint 8 — Analytics + Admin Panel  (Steps 45–53)')
para('Follow Steps 45–53 after completing Step 10 (Create Store). '
     'Admin endpoints require a staff/superuser account — see Step 45b.', size=11)
tip('No new migrations for this sprint. All analytics data is aggregated from existing tables.')

para('Step 45b: Create a Staff User (one-time dev setup)', bold=True)
code('docker compose exec django python manage.py shell -c "\nfrom apps.auth_app.models import User\nu, _ = User.objects.get_or_create(phone_number=\'+919000000001\', defaults={\'role\':\'admin\',\'full_name\':\'Admin\'})\nu.is_staff = True; u.is_superuser = True; u.save()\nprint(\'Done\')\n"')
para('Then run /auth/otp/send/ + /auth/otp/verify/ with +919000000001 to get admin_token.', size=10)

h2('STEP 45 — Vendor Dashboard  (Analytics)')
endpoint_header('GET', '{{base_url}}/analytics/vendor/', 'Vendor JWT', 'Sprint 8')
para('No body.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "store": {\n    "id": "...", "name": "Chennai Silk House", "category": "fashion",\n    "is_active": true, "is_verified": true, "is_open": true,\n    "follower_count": 2, "review_count": 2, "avg_rating": 4.5\n  },\n  "wallet": { "balance": "501.00" },\n  "subscription": { "plan": "Basic Plan", "expires_at": "2026-06-14T...", "is_active": true, "days_left": 29 },\n  "current_plan": { "name": "basic", "display_name": "Basic Plan", "video_limit": 20, "product_limit": 50 },\n  "products": { "total": 5, "active": 3, "draft": 1, "inactive": 1 },\n  "videos": { "total": 8, "ready": 3, "processing": 2, "pending": 3, "total_likes": 0, "total_views": 6 }\n}')
success('subscription is null if the vendor has never subscribed. current_plan defaults to Free.')

h2('STEP 46 — Video Stats  (Analytics)')
endpoint_header('GET', '{{base_url}}/analytics/vendor/videos/', 'Vendor JWT', 'Sprint 8')
para('No body.')
para('Expected Response  200 OK:', bold=True, size=10)
code('[\n  {\n    "id": "...", "title": "Summer Kurta Collection",\n    "status": "ready", "view_count": 4, "like_count": 0,\n    "duration_seconds": 45, "created_at": "2026-05-15T..."\n  }\n]')
tip('Ordered by most views first. All videos returned — not just ready ones.')

h2('STEP 47 — Product Stats  (Analytics)')
endpoint_header('GET', '{{base_url}}/analytics/vendor/products/', 'Vendor JWT', 'Sprint 8')
para('No body.')
para('Expected Response  200 OK:', bold=True, size=10)
code('[\n  {\n    "id": "...", "name": "Kanchipuram Silk Saree",\n    "status": "active", "base_price": "4500.00",\n    "wishlist_count": 1, "created_at": "2026-05-15T..."\n  }\n]')

h2('STEP 48 — Platform Stats  (Admin)')
endpoint_header('GET', '{{base_url}}/admin-panel/stats/', 'Staff JWT', 'Sprint 8')
para('No body. Use admin_token saved in Step 45b.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "users":    { "total": 7, "vendors": 2, "customers": 4, "active": 7 },\n  "stores":   { "total": 1, "active": 1, "verified": 1, "open": 1 },\n  "videos":   { "total": 8, "ready": 3, "total_views": 6, "total_likes": 0 },\n  "products": { "active": 1 },\n  "revenue":  { "subscription_revenue": "499.00", "total_topups": "1000.00" }\n}')

h2('STEP 49 — List All Stores  (Admin)')
endpoint_header('GET', '{{base_url}}/admin-panel/stores/', 'Staff JWT', 'Sprint 8')
para('No body. Optional query params: ?search=, ?is_active=, ?is_verified=, ?category=')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "count": 1,\n  "results": [\n    {\n      "id": "...", "name": "Chennai Silk House", "category": "fashion",\n      "address": "123 T Nagar", "locality": "T Nagar",\n      "is_active": true, "is_verified": true, "is_open": true,\n      "wallet_balance": "501.00",\n      "owner_phone": "+919999999999", "owner_name": "",\n      "product_count": 1, "video_count": 8, "created_at": "..."\n    }\n  ]\n}')
tip('Try: /admin-panel/stores/?is_verified=false to find stores pending verification.')

h2('STEP 50 — Verify a Store  (Admin)')
endpoint_header('PATCH', '{{base_url}}/admin-panel/stores/{{store_id}}/', 'Staff JWT', 'Sprint 8')
para('Body:', bold=True, size=10)
code('{ "is_verified": true }')
para('Expected Response  200 OK:', bold=True, size=10)
code('{ "id": "...", "is_verified": true, ... }')
success('You can also set: {"is_active": false} to deactivate, or {"is_open": true} to mark store as open.')

h2('STEP 51 — List All Users  (Admin)')
endpoint_header('GET', '{{base_url}}/admin-panel/users/', 'Staff JWT', 'Sprint 8')
para('No body. Optional query params: ?search=, ?role=vendor/customer, ?is_active=')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "count": 7,\n  "results": [\n    {\n      "id": "...", "phone_number": "+919999999999",\n      "role": "vendor", "is_active": true, "is_staff": false,\n      "store_name": "Chennai Silk House", "created_at": "..."\n    }\n  ]\n}')
tip('Vendors show store_name. Customers show null for store_name.')

h2('STEP 52 — Toggle User Active  (Admin)')
endpoint_header('POST', '{{base_url}}/admin-panel/users/{{customer_id}}/toggle-active/', 'Staff JWT', 'Sprint 8')
para('No body.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{ "message": "User deactivated successfully.", "user_id": "...", "is_active": false }')
para('Call again to re-activate:', size=10)
code('{ "message": "User activated successfully.", "user_id": "...", "is_active": true }')
error_box('Cannot toggle your own account — returns 400 with "Cannot deactivate your own account."')

h2('STEP 53 — Search Filter Tests  (Admin)')
para('Test the filter params on both admin endpoints:')
code('# Filter stores by verification status\nGET /admin-panel/stores/?is_verified=false\n\n# Filter stores by category\nGET /admin-panel/stores/?category=fashion\n\n# Search stores by name or owner phone\nGET /admin-panel/stores/?search=chennai\n\n# Filter users by role\nGET /admin-panel/users/?role=vendor\n\n# Search users by phone\nGET /admin-panel/users/?search=9999')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 9 — RESERVATIONS
# ════════════════════════════════════════════════════════════════
h1('Sprint 9 — Reservations  (Steps 54–58)')
para('Customers hold a product for 2 hours. Vendor confirms or rejects. Celery expires stale holds.', size=11)
tip('Hold duration: 2h by default. Controlled via RESERVATION_HOLD_HOURS in .env')

h2('STEP 54 — Create Reservation  (Customer)')
endpoint_header('POST', '{{base_url}}/reservations/', 'Bearer JWT', 'Sprint 9')
para('Body:', bold=True, size=10)
code('{\n  "store_id":   "{{store_id}}",\n  "product_id": "{{product_id}}",\n  "quantity":   2,\n  "note":       "Please keep ready by 6 PM"\n}')
para('Auto-save Tests script:', bold=True, size=10)
code('const r = pm.response.json();\nif (r.id) { pm.environment.set("reservation_id", r.id); }')
para('Expected Response  201 Created:', bold=True, size=10)
code('{\n  "id": "...",\n  "store":    { "id": "...", "name": "Chennai Silk House", "locality": "...", "phone": "..." },\n  "customer": { "id": "...", "phone_number": "+916000000001", "full_name": "" },\n  "product":  { "id": "...", "name": "Kanchipuram Silk Saree", "base_price": "4500.00" },\n  "quantity": 2,\n  "note": "Please keep ready by 6 PM",\n  "vendor_note": "",\n  "status": "pending",\n  "expires_at": "2026-05-15T23:25:49Z",\n  "hours_left": 2.0,\n  "created_at": "..."\n}')
success('hours_left = 2.0 on fresh creation. Decreases as time passes. 0 when expired.')

h2('STEP 55 — List Reservations')
endpoint_header('GET', '{{base_url}}/reservations/list/', 'Bearer JWT', 'Sprint 9')
para('No body. Behavior differs by role:')
para('  • Customer token → returns only their own reservations', size=10)
para('  • Vendor token → returns all reservations received by their store', size=10)
para('Expected Response  200 OK:', bold=True, size=10)
code('[\n  { "id": "...", "status": "pending", "quantity": 2, "product": { "name": "..." }, ... }\n]')

h2('STEP 56 — Vendor Confirms Reservation')
endpoint_header('PATCH', '{{base_url}}/reservations/{{reservation_id}}/status/', 'Vendor JWT', 'Sprint 9')
para('Body:', bold=True, size=10)
code('{ "status": "confirmed", "vendor_note": "Ready for pickup after 5 PM!" }')
para('Expected Response  200 OK:', bold=True, size=10)
code('{ "status": "confirmed", "vendor_note": "Ready for pickup after 5 PM!", ... }')
tip('Vendor can also send: {"status": "cancelled", "vendor_note": "Out of stock, sorry."} to reject.')

h2('STEP 57 — Vendor Marks Completed')
endpoint_header('PATCH', '{{base_url}}/reservations/{{reservation_id}}/status/', 'Vendor JWT', 'Sprint 9')
para('Body:', bold=True, size=10)
code('{ "status": "completed" }')
para('Expected Response  200 OK:', bold=True, size=10)
code('{ "status": "completed", ... }')
error_box('Must be confirmed first. Trying to complete a pending → 400 — Only confirmed reservations can be completed.')

h2('STEP 58 — Customer Cancels Reservation')
endpoint_header('POST', '{{base_url}}/reservations/{{reservation_id}}/cancel/', 'Bearer JWT', 'Sprint 9')
para('No body. Customer can only cancel their own pending reservation.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{ "status": "cancelled", ... }')
error_box('Cannot cancel a confirmed or completed reservation — returns 400.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 10 — GROUPS
# ════════════════════════════════════════════════════════════════
h1('Sprint 10 — Groups  (Steps 59–69)')
para('Collaborative shopping groups. Customers add friends by Profile ID. Vendors create exclusive groups for store followers. Phone numbers are never exposed.', size=11)
tip('Profile ID format: NK-XXXXXXXX (e.g. NK-A3X9K2). Users see their own profile_id in GET /auth/me/.')

h2('STEP 59 — Get Your Profile ID')
endpoint_header('GET', '{{base_url}}/auth/me/', 'Bearer JWT', 'Sprint 10')
para('No body. Returns your profile including profile_id.')
para('Expected Response  200 OK (new field):', bold=True, size=10)
code('{\n  "id": "...",\n  "profile_id": "NK-A3X9K2",\n  "phone_number": "+919999999999",\n  "role": "customer",\n  "full_name": "Rahul Kumar",\n  ...\n}')
tip('Share your profile_id with friends so they can add you to a group. It is private — only share with people you trust.')

h2('STEP 60 — Search User by Profile ID')
endpoint_header('GET', '{{base_url}}/auth/users/search/?profile_id=NK-A3X9K2', 'Bearer JWT', 'Sprint 10')
para('No body. Returns name and profile_id only — phone is never exposed.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "id": "...",\n  "profile_id": "NK-A3X9K2",\n  "full_name": "Priya Sharma"\n}')
success('Use this to confirm you have the right person before adding them to a group.')

h2('STEP 61 — Create Customer Group')
endpoint_header('POST', '{{base_url}}/groups/', 'Bearer JWT', 'Sprint 10')
para('Body:', bold=True, size=10)
code('{\n  "name": "Weekend Shopping",\n  "group_type": "customer"\n}')
para('Auto-save Tests script:', bold=True, size=10)
code('const r = pm.response.json();\nif (r.id) { pm.environment.set("group_id", r.id); }')
para('Expected Response  201 Created:', bold=True, size=10)
code('{\n  "id": "...",\n  "name": "Weekend Shopping",\n  "group_type": "customer",\n  "member_count": 1,\n  "created_by_name": "Rahul Kumar",\n  "created_by_profile_id": "NK-A3X9K2",\n  "store_name": null,\n  "is_active": true,\n  "created_at": "..."\n}')

h2('STEP 62 — Create Vendor Group')
endpoint_header('POST', '{{base_url}}/groups/', 'Vendor JWT', 'Sprint 10')
para('Body:', bold=True, size=10)
code('{\n  "name": "VIP Followers Deal",\n  "group_type": "vendor"\n}')
para('Expected Response  201 Created:', bold=True, size=10)
code('{\n  "id": "...",\n  "group_type": "vendor",\n  "store_name": "Chennai Silk House",\n  ...\n}')
error_box('Customer token with group_type=vendor → 400 — Only vendors can create vendor groups.')

h2('STEP 63 — Add Member to Customer Group (by Profile ID)')
endpoint_header('POST', '{{base_url}}/groups/{{group_id}}/members/add/', 'Admin JWT', 'Sprint 10')
para('Body:', bold=True, size=10)
code('{\n  "profile_id": "NK-B7Y3Z1"\n}')
para('Expected Response  201:', bold=True, size=10)
code('{ "message": "Priya Sharma added to group." }')
error_box('Wrong profile_id → 404 — No user found with this Profile ID.')

h2('STEP 64 — Eligible Members + Add to Vendor Group')
endpoint_header('GET', '{{base_url}}/groups/{{group_id}}/eligible-members/', 'Vendor Admin JWT', 'Sprint 10')
para('No body. Returns store followers not yet in the group.')
para('Expected Response  200 OK:', bold=True, size=10)
code('[\n  { "user_id": "...", "profile_id": "NK-C4X8W2", "full_name": "Anita Patel" },\n  ...\n]')
para('Then add by user_id:', bold=True, size=10)
code('POST /groups/{{group_id}}/members/add/\n{\n  "user_id": "uuid-from-eligible-list"\n}')
error_box('Non-follower user_id → 403 — User does not follow this store.')

h2('STEP 65 — Make Admin / Remove Admin')
endpoint_header('POST', '{{base_url}}/groups/{{group_id}}/members/{{user_id}}/make-admin/', 'Admin JWT', 'Sprint 10')
para('No body. Promotes a member to admin. Groups can have multiple admins.')
para('Expected Response  200:', bold=True, size=10)
code('{ "message": "Priya Sharma is now an admin." }')
para('To demote:', bold=True, size=10)
code('POST /groups/{{group_id}}/members/{{user_id}}/remove-admin/')
code('{ "message": "Priya Sharma is no longer an admin." }')
error_box('Cannot demote the group creator → 400 — Cannot remove admin role from the group creator.')

h2('STEP 66 — Share a Product in Group')
endpoint_header('POST', '{{base_url}}/groups/{{group_id}}/products/', 'Member JWT', 'Sprint 10')
para('Body:', bold=True, size=10)
code('{\n  "product_id": "{{product_id}}",\n  "note": "This looks perfect for the wedding!"\n}')
para('Auto-save Tests script:', bold=True, size=10)
code('const r = pm.response.json();\nif (r.id) { pm.environment.set("sp_id", r.id); }')
para('Expected Response  201 Created:', bold=True, size=10)
code('{\n  "id": "...",\n  "product_name": "Kanchipuram Silk Saree",\n  "product_price": "4500.00",\n  "store_name": "Chennai Silk House",\n  "note": "This looks perfect for the wedding!",\n  "is_finalized": false,\n  "shared_by_name": "Rahul Kumar",\n  "shared_by_profile_id": "NK-A3X9K2",\n  "finalized_by_name": null\n}')
error_box('Note with external URL → 400 — External links are not allowed. Only NearKart app links are permitted.')

h2('STEP 67 — List Shared Products')
endpoint_header('GET', '{{base_url}}/groups/{{group_id}}/products/', 'Member JWT', 'Sprint 10')
para('No body. Finalized products appear first.')
code('[\n  { "is_finalized": true, "product_name": "...", "finalized_by_name": "Rahul Kumar", ... },\n  { "is_finalized": false, "product_name": "...", ... }\n]')

h2('STEP 68 — Finalize a Product (Admin)')
endpoint_header('POST', '{{base_url}}/groups/{{group_id}}/products/{{sp_id}}/finalize/', 'Admin JWT', 'Sprint 10')
para('No body. Admin marks this shared product as the group\'s final choice.')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "is_finalized": true,\n  "finalized_by_name": "Rahul Kumar",\n  ...\n}')
error_box('Already finalized → 400 — Product is already finalized.')

h2('STEP 69 — Leave / Delete Group')
endpoint_header('POST', '{{base_url}}/groups/{{group_id}}/leave/', 'Member JWT', 'Sprint 10')
para('No body. Any non-creator member can leave.')
code('{ "message": "You have left the group." }')
para('Delete (creator only):', bold=True, size=10)
endpoint_header('DELETE', '{{base_url}}/groups/{{group_id}}/', 'Creator JWT', 'Sprint 10')
code('{ "message": "Group deleted." }')
error_box('Creator trying to leave → 400 — Group creator cannot leave. Delete the group instead.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 11 — NOTIFICATIONS
# ════════════════════════════════════════════════════════════════
h1('Sprint 11 — Notifications  (Steps 70–74)')
para('In-app inbox + Firebase push notifications. Every key event (message, reservation, follow, review, like, wallet top-up, group action) sends a notification to the relevant user.', size=11)
tip('In dev mode: push notifications are LOGGED to console, not sent via Firebase. Look for [FCM-DEV] lines in Django logs.')

h2('STEP 70 — Register Device FCM Token')
para('Must be called after login to receive push notifications on the device.', bold=True, size=10)
endpoint_header('POST', '{{base_url}}/notifications/device-token/', 'Bearer JWT', 'Sprint 11')
code('{\n  "fcm_token": "ExponentPushToken[test_token_abc123]",\n  "device_type": "android"\n}')
para('Expected response:')
code('{ "message": "Device token registered." }')
tip('device_type choices: android / ios / web')

h2('STEP 71 — List Notifications (Inbox)')
endpoint_header('GET', '{{base_url}}/notifications/', 'Bearer JWT', 'Sprint 11')
para('No body required. Returns last 50 notifications for the authenticated user.')
para('Expected response:')
code('[\n  {\n    "id": "<uuid>",\n    "notification_type": "new_message",\n    "title": "New message from Ravi Fashion",\n    "body": "You have a new message.",\n    "data": {"conversation_id": "<uuid>"},\n    "is_read": false,\n    "created_at": "2026-05-15T09:00:00Z"\n  }\n]')
note('Tests script to save notification_id:\nconst data = pm.response.json();\nif (data.length > 0) pm.collectionVariables.set("notification_id", data[0].id);',
     'FFF3CD', (0x85, 0x62, 0x04))

h2('STEP 72 — Get Unread Count (Badge)')
endpoint_header('GET', '{{base_url}}/notifications/unread-count/', 'Bearer JWT', 'Sprint 11')
para('No body required. Used to display notification badge on the app icon.')
code('{ "unread_count": 3 }')

h2('STEP 73 — Mark One Notification Read')
endpoint_header('POST', '{{base_url}}/notifications/{{notification_id}}/read/', 'Bearer JWT', 'Sprint 11')
para('No body required. Marks a single notification as read.')
code('{ "message": "Marked as read." }')
tip('Calling this multiple times is safe — idempotent.')

h2('STEP 74 — Mark All Notifications Read')
endpoint_header('POST', '{{base_url}}/notifications/read-all/', 'Bearer JWT', 'Sprint 11')
para('No body required. Marks all of the user\'s unread notifications as read.')
code('{ "marked_read": 3 }')
para('After this, GET /notifications/unread-count/ returns 0.')

h2('Trigger Notifications (side-effects from other endpoints)')
table(
    ['Trigger Action', 'Who Gets Notified', 'Notification Type'],
    [
        ('Customer sends WS message',           'Other conversation party',     'new_message'),
        ('Customer creates reservation',        'Vendor',                       'reservation_created'),
        ('Vendor confirms reservation',         'Customer',                     'reservation_confirmed'),
        ('Vendor cancels reservation',          'Customer',                     'reservation_cancelled'),
        ('Reservation hold expires (Celery)',   'Customer',                     'reservation_expired'),
        ('User follows store',                  'Vendor',                       'new_follower'),
        ('User reviews store',                  'Vendor',                       'new_review'),
        ('Vendor sets is_open=true',            'All followers (bulk)',          'store_opened'),
        ('User likes video',                    'Vendor',                       'video_liked'),
        ('Admin tops up wallet',                'Vendor',                       'wallet_topup'),
        ('Sub expiring in ~3 days (Celery)',    'Vendor',                       'subscription_expiring'),
        ('Sub expired in last 24h (Celery)',    'Vendor',                       'subscription_expired'),
        ('Admin adds member to group',          'Added user',                   'group_added'),
        ('Admin removes member from group',     'Removed user',                 'group_removed'),
        ('Member shares product in group',      'All other group members',      'group_product_shared'),
        ('Admin finalizes product in group',    'All group members',            'group_product_finalized'),
        ('Admin promotes member to admin',      'Promoted user',                'group_admin_promoted'),
    ],
    col_widths=[2.2, 1.8, 2.6],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SPRINT 12 — PRODUCTION + RAZORPAY PAYMENTS
# ════════════════════════════════════════════════════════════════
h1('Sprint 12 — Production + Razorpay Payments  (Steps 75–79)')
para('Sprint 12 completes two things: (1) Store Hours management API, and (2) Razorpay payment flow so vendors can pay for subscriptions through the app instead of needing admin top-ups.', size=11)
tip('Dev mode: All Razorpay calls use placeholder credentials — no real money moves.\n'
    'Replace RAZORPAY_KEY_ID / RAZORPAY_KEY_SECRET / RAZORPAY_WEBHOOK_SECRET in .env with real keys when going live.\n'
    'Get keys from: https://dashboard.razorpay.com/app/keys')

note('📌  Production Architecture Summary:\n'
     '  Internet → AWS ALB (HTTPS/SSL termination) → Nginx (port 80) → Gunicorn+Uvicorn (port 8000) → Django\n'
     '  Static files → S3 (via django-storages). Media (videos/images) → S3.\n'
     '  Logs → JSON structured to stdout → CloudWatch. Errors → Sentry.\n'
     '  CI/CD: lint → test → Docker build → ECR → auto-deploy staging → manual approval → production.',
     'D1ECF1', (0x0C, 0x54, 0x60))

# ── STEP 75 ──
divider()
h2('STEP 75 — Get Store Hours')
endpoint_header('GET', '{{base_url}}/stores/{{store_id}}/hours/', 'Store Owner JWT', 'Sprint 12')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('No body. Returns operating hours for the vendor\'s store.')
para('Expected Response  200 OK:', bold=True, size=10)
code('[\n  {"id": "...", "day": 0, "open_time": "09:00", "close_time": "21:00", "is_closed": false},\n  {"id": "...", "day": 1, "open_time": "09:00", "close_time": "21:00", "is_closed": false},\n  {"id": "...", "day": 6, "open_time": "00:00", "close_time": "00:00", "is_closed": true}\n]')
tip('day field: 0=Monday, 1=Tuesday, 2=Wednesday, 3=Thursday, 4=Friday, 5=Saturday, 6=Sunday\n'
    'Returns [] if no hours have been set yet.\n'
    'Only the store owner can view or set hours — 403 for any other token.')

# ── STEP 76 ──
divider()
h2('STEP 76 — Set Store Hours  (Replace All)')
endpoint_header('PUT', '{{base_url}}/stores/{{store_id}}/hours/', 'Store Owner JWT', 'Sprint 12')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body  (array of up to 7 day entries):', bold=True, size=10)
code('[\n  {"day": 0, "open_time": "09:00", "close_time": "21:00", "is_closed": false},\n  {"day": 1, "open_time": "09:00", "close_time": "21:00", "is_closed": false},\n  {"day": 2, "open_time": "09:00", "close_time": "21:00", "is_closed": false},\n  {"day": 3, "open_time": "09:00", "close_time": "21:00", "is_closed": false},\n  {"day": 4, "open_time": "09:00", "close_time": "21:00", "is_closed": false},\n  {"day": 5, "open_time": "10:00", "close_time": "22:00", "is_closed": false},\n  {"day": 6, "open_time": "00:00", "close_time": "00:00", "is_closed": true}\n]')
para('Expected Response  200 OK:', bold=True, size=10)
code('[\n  {"id": "...", "day": 0, "open_time": "09:00", "close_time": "21:00", "is_closed": false},\n  ...\n]')
warning('This REPLACES all existing hours — omitted days are deleted.\n'
        'To close a day: set is_closed=true (open_time and close_time are ignored for closed days).\n'
        'Cannot have duplicate day entries — 400 if you send day 0 twice.')
success('After setting hours, the store detail endpoint (GET /stores/<id>/) returns hours in the response.')

# ── STEP 77 ──
divider()
h2('STEP 77 — Initiate Razorpay Payment  (Vendor Only)')
endpoint_header('POST', '{{base_url}}/billing/payment/initiate/', 'Vendor JWT', 'Sprint 12')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body:', bold=True, size=10)
code('{\n  "plan_name": "basic"\n}')
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "order_id":        "order_DEV_store_abc12345",\n  "amount":          49900,\n  "currency":        "INR",\n  "plan_name":       "basic",\n  "receipt":         "store_abc12345_basic_1716000000",\n  "razorpay_key_id": "rzp_test_PLACEHOLDER"\n}')
note('amount is in paise (₹ × 100):\n  Basic plan ₹499 → amount: 49900\n  Premium plan ₹999 → amount: 99900',
     'D1ECF1', (0x0C, 0x54, 0x60))
tip('Dev mode: order_id starts with "order_DEV_" — no real Razorpay call.\n'
    'Production: use order_id + razorpay_key_id to open the Razorpay checkout SDK in the mobile app.\n'
    'After the user pays in the SDK, the SDK returns razorpay_payment_id and razorpay_signature.\n'
    'Send those to STEP 78 to complete the payment.')
para('Auto-save Tests script:', bold=True, size=10)
code('const r = pm.response.json();\nif (r.order_id) {\n    pm.environment.set("razorpay_order_id", r.order_id);\n    console.log("Order ID saved:", r.order_id);\n}')
table(
    ['plan_name', 'Price', 'amount (paise)'],
    [
        ('basic',   '₹499/month',  '49900'),
        ('premium', '₹999/month', '99900'),
        ('free',    '₹0',          '400 error — use POST /billing/subscribe/ instead'),
    ],
    col_widths=[1.2, 1.4, 4.0],
)
error_box('400 — free plan needs no payment — use POST /billing/subscribe/ with plan_name: "free" directly\n'
          '     404 — plan not found — only valid paid plans: basic | premium\n'
          '     403 — vendor access only — use vendor token\n'
          '     404 — you do not have a store yet — create a store first')

# ── STEP 78 ──
divider()
h2('STEP 78 — Verify Payment and Activate Subscription  (Vendor Only)')
endpoint_header('POST', '{{base_url}}/billing/payment/verify/', 'Vendor JWT', 'Sprint 12')
para('Authorization: Bearer {{vendor_token}}', size=10, indent=0.2)
para('Content-Type: application/json', size=10, indent=0.2)
para('Request Body:', bold=True, size=10)
code('{\n  "razorpay_order_id":   "{{razorpay_order_id}}",\n  "razorpay_payment_id": "pay_DEV_test12345",\n  "razorpay_signature":  "mock_signature_dev",\n  "plan_name":           "basic"\n}')
note('In dev mode: razorpay_payment_id and razorpay_signature can be any non-empty string.\n'
     'In production: these come from the Razorpay checkout SDK callback.',
     'FFF3CD', (0x85, 0x62, 0x04))
para('Expected Response  200 OK:', bold=True, size=10)
code('{\n  "id": "...",\n  "store_name": "Fashion Hub",\n  "plan": {\n    "name": "basic",\n    "display_name": "Basic Plan",\n    "price": "499.00",\n    "video_limit": 20,\n    "product_limit": 50\n  },\n  "started_at": "2026-05-15T10:00:00Z",\n  "expires_at": "2026-06-14T10:00:00Z",\n  "is_active": true,\n  "days_left": 29\n}')
success('On success:\n'
        '  1. ₹499 (or ₹999 for premium) is credited to the vendor\'s wallet\n'
        '  2. Subscription is immediately activated\n'
        '  3. Wallet transaction with reference_id = razorpay_payment_id is recorded\n'
        '  4. wallet_topup notification is sent to the vendor')
para('After verify — confirm wallet + transactions:', bold=True, size=10)
code('GET {{base_url}}/billing/wallet/        → wallet_balance shows the net\nGET {{base_url}}/billing/transactions/   → 2 records: topup (+499) + subscription (-499)')
error_box('400 — all fields required (any of the 4 fields is missing or empty)\n'
          '     400 — payment_failed (production only — HMAC signature mismatch)\n'
          '          Do NOT retry a failed payment — contact support\n'
          '     404 — plan not found (plan_name does not match a real active plan)\n'
          '     400 — subscription_failed (topup OK but subscribe errored — check logs)')

# ── STEP 79 ──
divider()
h2('STEP 79 — Webhook Simulation  (Dev Mode)')
endpoint_header('POST', '{{base_url}}/billing/payment/webhook/', 'None (signature only)', 'Sprint 12')
para('No JWT. Razorpay calls this endpoint automatically after payment.captured event.', size=10)
para('Content-Type: application/json', size=10, indent=0.2)
para('Header: X-Razorpay-Signature: mock_sig   (any value in dev mode)', size=10, indent=0.2)
para('Request Body (simulates Razorpay webhook payload):', bold=True, size=10)
code('{\n  "event": "payment.captured",\n  "payload": {\n    "payment": {\n      "entity": {\n        "id":       "pay_DEV_webhook001",\n        "order_id": "order_DEV_store_abc",\n        "notes": {\n          "store_id": "{{store_id}}",\n          "plan":     "basic"\n        }\n      }\n    }\n  }\n}')
para('Expected Response  200 OK:', bold=True, size=10)
code('{ "status": "ok" }')
success('Wallet funded + subscription activated (if not already done by verify).')
para('Test idempotency (call same webhook twice):', bold=True, size=10)
code('Send exact same request again → { "status": "already_processed" }\nNo duplicate transaction is created — the webhook checks reference_id first.')
tip('Register this URL in Razorpay Dashboard → Settings → Webhooks (production):\n'
    '  https://api.nearkart.in/api/v1/billing/payment/webhook/\n'
    '  Events to subscribe: payment.captured\n\n'
    'The webhook is the safety net — if the app crashes after payment but before verify, the webhook activates the subscription automatically.')
warning('Production: Razorpay sends a real HMAC-SHA256 signature in X-Razorpay-Signature header.\n'
        'Set RAZORPAY_WEBHOOK_SECRET in .env to the secret shown in Razorpay Dashboard → Webhooks.\n'
        'Without the correct secret, webhook requests return 400 — invalid_signature.')

h2('Full Razorpay Payment Flow  (End to End)')
table(
    ['Step', 'Who', 'Action', 'Result'],
    [
        ('1', 'Vendor',   'GET /billing/plans/',                       'See plan prices — basic ₹499, premium ₹999'),
        ('2', 'Vendor',   'POST /billing/payment/initiate/ {plan_name: "basic"}', 'Get order_id + razorpay_key_id'),
        ('3', 'App',      'Open Razorpay checkout SDK',               'User enters card/UPI → payment processed'),
        ('4', 'SDK',      'Returns razorpay_payment_id + razorpay_signature', 'Callback to the app'),
        ('5', 'Vendor',   'POST /billing/payment/verify/',            'Signature verified → wallet funded → subscription active'),
        ('6', 'Razorpay', 'POST /billing/payment/webhook/ (auto)',    'Backup — fires even if step 5 failed'),
        ('7', 'Vendor',   'GET /billing/subscription/',               'Confirm is_active=true and plan is set'),
    ],
    col_widths=[0.4, 0.8, 2.8, 2.6],
)

h2('Production Go-Live Checklist  (Razorpay)')
bullet('Get live keys from https://dashboard.razorpay.com/app/keys')
bullet('Set RAZORPAY_KEY_ID=rzp_live_... in .env.production')
bullet('Set RAZORPAY_KEY_SECRET=... in .env.production')
bullet('Register webhook in Razorpay Dashboard → Settings → Webhooks')
bullet('  URL: https://api.nearkart.in/api/v1/billing/payment/webhook/')
bullet('  Events: payment.captured')
bullet('Copy the webhook secret → set RAZORPAY_WEBHOOK_SECRET in .env.production')
bullet('Test with a ₹1 real transaction on staging before going live')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 4 — ERROR TESTING
# ════════════════════════════════════════════════════════════════
h1('4. Error Cases to Test')

h2('Phone Number Error Cases (Auth)')
table(
    ['Wrong Input', 'Error Response', 'Fix'],
    [
        ('{"phone_number": "9999999999"}',     '400 — Enter a valid Indian mobile number in +91XXXXXXXXXX format.', 'Add +91 prefix: +919999999999'),
        ('{"phone_number": "919999999999"}',   '400 — Enter a valid Indian mobile number in +91XXXXXXXXXX format.', 'Add + sign: +919999999999'),
        ('{"phone_number": "+91 9999999999"}', '400 — Enter a valid Indian mobile number in +91XXXXXXXXXX format.', 'Remove space: +919999999999'),
        ('{"phone_number": "+915000000000"}',  '400 — Enter a valid Indian mobile number in +91XXXXXXXXXX format.', 'First digit must be 6–9: +916000000000'),
        ('{"phone_number": ""}',               '400 — This field may not be blank.',                                'Provide a phone number'),
        ('{}  (no phone_number key)',           '400 — This field is required.',                                    'Add phone_number field to request body'),
    ],
    col_widths=[2.3, 2.2, 2.1],
)

h2('OTP Error Cases')
table(
    ['Test', 'Request / Situation', 'Expected Response'],
    [
        ('Wrong OTP (1st attempt)',   '/otp/verify/ with otp: "000000"',            '400 — Invalid OTP. 4 attempt(s) remaining.'),
        ('Wrong OTP (5th attempt)',   '/otp/verify/ with wrong OTP 5 times',        '400 — OTP session locked. Request a new OTP.'),
        ('Expired OTP (10 min old)', '/otp/verify/ after 10 minutes',               '400 — OTP expired. Request a new one.'),
        ('OTP not sent yet',         '/otp/verify/ without calling /otp/send/ first','400 — No active OTP found for this number.'),
        ('Too many OTP sends',       'POST /otp/send/ more than 5 times / 5 min',   '429 — Request was throttled. Try again later.'),
    ],
    col_widths=[1.6, 2.6, 2.4],
)

h2('Token / Auth Error Cases')
table(
    ['Test', 'Request', 'Expected Response'],
    [
        ('No Authorization header',    'GET /auth/me/ with no header',                          '401 — authentication_failed'),
        ('Wrong format (no Bearer)',   'Authorization: eyJhbGci... (missing Bearer)',            '401 — authentication_failed'),
        ('Expired access token',       'GET /auth/me/ with token older than 1 hour',            '401 — token_invalid'),
        ('Blacklisted refresh token',  'POST /token/refresh/ after calling /logout/',           '401 — token_invalid'),
        ('Invalid coordinates',        'PUT /auth/me/location/ → latitude: 200',               '400 — validation_error'),
    ],
    col_widths=[1.8, 2.6, 2.2],
)

h2('Store Error Cases')
table(
    ['Test', 'Request', 'Expected Response'],
    [
        ('Customer creates store',     'POST /stores/ with customer token',                 '403 — Vendor access only'),
        ('Vendor creates 2nd store',   'POST /stores/ again with same vendor',              '400 — You already have a store'),
        ('Wrong owner updates store',  'PUT /stores/<id>/update/ with different vendor',    '403 — permission_denied'),
        ('Missing lat/lng',            'GET /stores/nearby/ without ?lat=&lng=',            '400 — lat and lng are required numbers'),
        ('Invalid store UUID',         'GET /stores/00000000-0000-0000-0000-000000000000/', '404 — not_found'),
        ('Store is_active=false',      'GET /stores/<id>/ after deactivating store',        '404 — not_found'),
        ('Wrong phone in body',        'POST /stores/ with phone: "9876543210"',            '400 — validation_error (phone format)'),
    ],
    col_widths=[1.8, 2.6, 2.2],
)

h2('Product Error Cases')
table(
    ['Test', 'Request', 'Expected Response'],
    [
        ('Vendor has no store yet',    'POST /products/ before creating store',              '400 — Create a store first'),
        ('Duplicate SKU',              'POST /products/ with sku already used by any store', '400 — validation_error on sku'),
        ('Get draft product',          'GET /products/<id>/ where status=draft',             '404 — not_found'),
        ('Get invisible product',      'GET /products/<id>/ where is_visible=false',         '404 — not_found'),
        ('Missing q param',            'GET /products/search/ without ?q=',                  '400 — Search query q is required'),
        ('Customer deletes product',   'DELETE /products/<id>/update/ with customer token',  '403 — permission_denied'),
    ],
    col_widths=[1.8, 2.6, 2.2],
)

h2('Chat Error Cases  (Sprint 5)')
table(
    ['Test', 'Request / Situation', 'Expected Response'],
    [
        ('Missing store_id',           'POST /conversations/start/ with empty body',         '400 — store_id is required'),
        ('Invalid store',              'POST /conversations/start/ — non-existent store_id', '404 — Store not found'),
        ('No auth on list',            'GET /conversations/ — no Authorization header',       '401 — authentication_failed'),
        ('Outsider views messages',    'GET /conversations/<id>/messages/ — not a member',    '403 — permission_denied'),
        ('Outsider marks read',        'PATCH /conversations/<id>/read/ — not a member',      '403 — permission_denied'),
        ('WS no token',                'Connect without ?token= in URL',                      'WS close code 4001'),
        ('WS wrong token',             'Connect with expired/invalid JWT token',              'WS close code 4001'),
        ('WS wrong conversation',      'Connect to conversation user is not part of',         'WS close code 4003'),
        ('WS empty message',           'Send {"type": "chat_message", "content": ""}',        'Ignored silently'),
        ('WS unknown message type',    'Send {"type": "ping"}',                               'Ignored silently'),
    ],
    col_widths=[2.0, 2.5, 2.1],
)

h2('Blacklist Error Cases  (Sprint 6)')
table(
    ['Test', 'Request / Situation', 'Expected Response'],
    [
        ('Block — not store owner',       'POST /stores/<id>/blacklist/<cust_id>/ — different vendor', '403 — permission_denied'),
        ('Block — customer not found',    'POST /stores/<id>/blacklist/<bad_uuid>/',                   '404 — Customer not found'),
        ('Block — store not found',       'POST /stores/<bad_uuid>/blacklist/<cust_id>/',              '404 — Store not found'),
        ('Block — with customer token',   'POST /stores/<id>/blacklist/<id>/ — customer token',        '403 — Vendor access only'),
        ('Block — no auth',               'POST /stores/<id>/blacklist/<id>/ — no header',             '401 — authentication_failed'),
        ('List — not store owner',        'GET /stores/<id>/blacklist/ — different vendor',            '403 — permission_denied'),
        ('Blocked follows store',         'POST /stores/<id>/follow/ — blocked customer token',        '403 — blacklisted'),
        ('Blocked reviews store',         'POST /stores/<id>/review/ — blocked customer token',        '403 — blacklisted'),
        ('Blocked starts conversation',   'POST /conversations/start/ — blocked customer token',       '403 — blacklisted'),
        ('Blocked connects via WS',       'ws://.../ws/conversations/<id>/?token=<blocked_token>',     'WS close code 4003'),
    ],
    col_widths=[1.9, 2.6, 2.1],
)

h2('Billing Error Cases  (Sprint 7)')
table(
    ['Test', 'Request / Situation', 'Expected Response'],
    [
        ('Top-up — zero amount',         'POST /billing/topup/ {"amount": "0"}',                '400 — amount must be a positive number'),
        ('Top-up — negative amount',     'POST /billing/topup/ {"amount": "-100"}',             '400 — validation_error'),
        ('Top-up — non-numeric',         'POST /billing/topup/ {"amount": "abc"}',              '400 — validation_error'),
        ('Top-up — no auth',             'POST /billing/topup/ — no Authorization header',      '401 — authentication_failed'),
        ('Top-up — customer token',      'POST /billing/topup/ — customer token',               '403 — Vendor access only'),
        ('Subscribe — bad plan name',    'POST /billing/subscribe/ {"plan_name": "gold"}',      '404 — Plan not found'),
        ('Subscribe — insufficient bal', 'POST /billing/subscribe/ {"plan_name": "premium"} — balance too low', '400 — insufficient_balance'),
        ('Subscription status — no sub', 'GET /billing/subscription/ — never subscribed',       '404 — No subscription found'),
        ('Wallet — no store',            'GET /billing/wallet/ — vendor with no store',         '404 — not_found'),
        ('Video limit reached',          'POST /videos/request-upload/ on Free plan (>3 videos)', '403 — plan_limit_reached'),
        ('Product limit reached',        'POST /products/ on Free plan (>10 products)',         '403 — plan_limit_reached'),
    ],
    col_widths=[1.9, 2.6, 2.1],
)

h2('Reservation Error Cases  (Sprint 9)')
table(
    ['Test', 'Request / Situation', 'Expected Response'],
    [
        ('Invalid product',          'POST /reservations/ — product_id not in store',        '404 — Product not found or not available'),
        ('Draft product',            'POST /reservations/ — product status=draft',            '404 — Product not found or not available'),
        ('Invalid store',            'POST /reservations/ — wrong store_id',                 '404 — Store not found'),
        ('Blacklisted customer',     'POST /reservations/ — blocked by store',               '403 — blacklisted'),
        ('No auth',                  'POST /reservations/ — no Authorization header',         '401 — authentication_failed'),
        ('Confirm already-confirmed','PATCH /status/ {"status":"confirmed"} — already done', '400 — Cannot confirmed a confirmed reservation'),
        ('Complete a pending',       'PATCH /status/ {"status":"completed"} — pending res',  '400 — Only confirmed reservations can be completed'),
        ('Cancel confirmed',         'POST /cancel/ — reservation is confirmed',             '400 — Cannot cancel a confirmed reservation'),
        ('Cancel other customer',    'POST /cancel/ — different customer token',             '404 — Reservation not found'),
        ('Vendor updates own res',   'PATCH /status/ with wrong vendor token',               '404 — Reservation not found'),
        ('Customer calls /status/',  'PATCH /status/ with customer token',                   '403 — Vendor access only'),
    ],
    col_widths=[1.9, 2.6, 2.1],
)

h2('Notifications Error Cases  (Sprint 11)')
table(
    ['Test', 'Request / Situation', 'Expected Response'],
    [
        ('No fcm_token',             'POST /device-token/ — body missing fcm_token',              '400 — This field is required'),
        ('Invalid device_type',      'POST /device-token/ {"device_type": "smartwatch"}',         '400 — Invalid device type'),
        ('Bad notification UUID',    'POST /notifications/<not-a-uuid>/read/',                    '400 — Not a valid UUID'),
        ('Wrong user notification',  'POST /notifications/<other-users-notif-id>/read/',          '404 — Not found (ownership enforced)'),
        ('No auth on inbox',         'GET /notifications/ — no Authorization header',             '401 — authentication_failed'),
        ('No auth on unread count',  'GET /notifications/unread-count/ — no header',              '401 — authentication_failed'),
        ('Read-all with no unread',  'POST /notifications/read-all/ — all already read',          '200 — {"marked_read": 0}'),
    ],
    col_widths=[1.9, 2.6, 2.1],
)

h2('Groups Error Cases  (Sprint 10)')
table(
    ['Test', 'Request / Situation', 'Expected Response'],
    [
        ('No profile_id param',       'GET /auth/users/search/ — no ?profile_id=',            '400 — profile_id query param is required'),
        ('Wrong profile_id',          'GET /auth/users/search/?profile_id=NK-ZZZZZZZZ',       '404 — No user found with this Profile ID'),
        ('Vendor type as customer',   'POST /groups/ {"group_type":"vendor"} — customer JWT', '400 — Only vendors can create vendor groups'),
        ('Vendor — no store',         'POST /groups/ {"group_type":"vendor"} — no store yet', '400 — Create a store first'),
        ('Non-admin adds member',     'POST /groups/<id>/members/add/ — member token',         '403 — Only group admin can add members'),
        ('Already a member',          'POST /groups/<id>/members/add/ — duplicate',            '400 — User is already a member of this group'),
        ('Non-follower vendor group', 'POST add/ {"user_id":"..."} — user not following store','403 — User does not follow this store'),
        ('Remove creator',            'DELETE /members/<creator_id>/remove/',                  '400 — Cannot remove the group creator'),
        ('Creator leaves',            'POST /groups/<id>/leave/ — creator token',              '400 — Group creator cannot leave'),
        ('External link in note',     'POST /products/ note: "http://youtube.com/xyz"',        '400 — External links are not allowed'),
        ('Non-admin finalizes',       'POST /products/<id>/finalize/ — member token',          '403 — Only group admin can finalize products'),
        ('Finalize twice',            'POST /products/<id>/finalize/ — already done',          '400 — Product is already finalized'),
        ('Non-admin eligible-members','GET /eligible-members/ — member token',                 '403 — Only group admin can view eligible members'),
        ('Eligible on customer group','GET /eligible-members/ on customer group',              '400 — Eligible members only available for vendor groups'),
        ('Delete — not creator',      'DELETE /groups/<id>/ — non-creator admin token',        '403 — Only the group creator can delete'),
    ],
    col_widths=[1.9, 2.6, 2.1],
)

h2('Store Hours Error Cases  (Sprint 12)')
table(
    ['Test', 'Request / Situation', 'Expected Response'],
    [
        ('GET hours — not owner',     'GET /stores/<id>/hours/ — different vendor token',  '403 — permission_denied'),
        ('GET hours — customer token','GET /stores/<id>/hours/ — customer token',          '403 — permission_denied'),
        ('PUT hours — duplicate day', 'PUT /stores/<id>/hours/ — two entries with day: 0', '400 — Duplicate day entries are not allowed'),
        ('PUT hours — not owner',     'PUT /stores/<id>/hours/ — different vendor token',  '403 — permission_denied'),
        ('PUT hours — invalid day',   'PUT /stores/<id>/hours/ — day: 7 (valid: 0–6)',     '400 — validation_error'),
        ('PUT hours — invalid time',  'PUT /stores/<id>/hours/ — open_time: "25:00"',      '400 — validation_error'),
    ],
    col_widths=[1.9, 2.6, 2.1],
)

h2('Razorpay Payment Error Cases  (Sprint 12)')
table(
    ['Test', 'Request / Situation', 'Expected Response'],
    [
        ('Initiate — free plan',         'POST /payment/initiate/ {"plan_name": "free"}',           '400 — free plan needs no payment'),
        ('Initiate — unknown plan',       'POST /payment/initiate/ {"plan_name": "gold"}',           '404 — plan not found'),
        ('Initiate — customer token',     'POST /payment/initiate/ — customer token',                '403 — Vendor access only'),
        ('Initiate — no store',           'POST /payment/initiate/ — vendor with no store',         '404 — you do not have a store yet'),
        ('Initiate — no auth',            'POST /payment/initiate/ — no Authorization header',      '401 — authentication_failed'),
        ('Verify — missing fields',       'POST /payment/verify/ — any field empty or missing',     '400 — all fields required'),
        ('Verify — wrong signature (prod)','POST /payment/verify/ — signature mismatch',            '400 — payment_failed'),
        ('Verify — unknown plan',         'POST /payment/verify/ — plan_name: "gold"',              '404 — plan not found'),
        ('Verify — no auth',              'POST /payment/verify/ — no Authorization header',        '401 — authentication_failed'),
        ('Webhook — wrong sig (prod)',    'POST /payment/webhook/ — wrong X-Razorpay-Signature',    '400 — invalid_signature'),
        ('Webhook — invalid JSON',        'POST /payment/webhook/ — non-JSON body',                  '400 — invalid_payload'),
        ('Webhook — duplicate payment',   'POST /payment/webhook/ — same payment_id twice',          '200 — already_processed (idempotent)'),
        ('Webhook — unknown event',       'POST /payment/webhook/ {"event": "payment.failed"}',     '200 — ok (ignored gracefully)'),
        ('Webhook — with JWT header',     'POST /payment/webhook/ with Authorization header',       'JWT ignored — endpoint has no auth'),
    ],
    col_widths=[1.9, 2.6, 2.1],
)

h2('Analytics Error Cases  (Sprint 8)')
table(
    ['Test', 'Request / Situation', 'Expected Response'],
    [
        ('Vendor Dashboard — no store', 'GET /analytics/vendor/ — vendor with no store',         '400 — Create a store first'),
        ('Dashboard — no auth',         'GET /analytics/vendor/ — no Authorization header',       '401 — authentication_failed'),
        ('Dashboard — customer token',  'GET /analytics/vendor/ — customer token',               '403 — Vendor access only'),
        ('Admin stats — vendor token',  'GET /admin-panel/stats/ — vendor token',                '403 — permission_denied'),
        ('Admin stats — no auth',       'GET /admin-panel/stats/ — no header',                   '401 — authentication_failed'),
        ('Store update — wrong UUID',   'PATCH /admin-panel/stores/<bad_uuid>/',                 '404 — Store not found'),
        ('User not found',              'POST /admin-panel/users/<bad_uuid>/toggle-active/',      '404 — User not found'),
        ('Toggle own account',          'POST /admin-panel/users/<own_id>/toggle-active/',        '400 — Cannot deactivate your own account'),
        ('Store list — vendor token',   'GET /admin-panel/stores/ — vendor token',               '403 — permission_denied'),
        ('User list — vendor token',    'GET /admin-panel/users/ — vendor token',                '403 — permission_denied'),
    ],
    col_widths=[1.9, 2.6, 2.1],
)

h2('Video Error Cases  (Sprint 4)')
table(
    ['Test', 'Request', 'Expected Response'],
    [
        ('No store — request upload',    'POST /videos/request-upload/ before store created',   '400 — Create a store first'),
        ('Missing title',                'POST /videos/request-upload/ with no title field',     '400 — This field is required'),
        ('Customer calls request-upload','POST /videos/request-upload/ with customer token',     '403 — Vendor access only'),
        ('Confirm already-ready video',  'POST /videos/<id>/confirm-upload/ on ready video',     '400 — Video already in "ready" state'),
        ('Duration too long',            'POST /videos/<id>/confirm-upload/ — duration_seconds=120', '400 — cannot exceed 60 seconds'),
        ('Confirm another vendor video', 'POST /videos/<id>/confirm-upload/ — wrong vendor',     '404 — not_found'),
        ('Feed without lat/lng',         'GET /videos/feed/ — no ?lat=&lng= params',             '400 — lat and lng are required numbers'),
        ('Detail on processing video',   'GET /videos/<id>/ while status=processing',            '404 — not_found (only ready shown)'),
        ('Like without auth',            'POST /videos/<id>/like/ — no auth header',             '401 — authentication_failed'),
        ('Update another vendor video',  'PATCH /videos/<id>/update/ — different vendor token',  '404 — not_found'),
        ('Delete another vendor video',  'DELETE /videos/<id>/delete/ — different vendor',       '404 — not_found'),
        ('Delete with customer token',   'DELETE /videos/<id>/delete/ — customer token',         '403 — Vendor access only'),
    ],
    col_widths=[2.0, 2.4, 2.2],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 5 — POSTMAN QUICK COLLECTION
# ════════════════════════════════════════════════════════════════
h1('5. Postman Quick Setup')

h2('Collection Folder Structure (recommended)')
code('NearKart Local\n├── 0. Setup\n│   └── Health Check\n├── 1. Auth\n│   ├── Send OTP (Vendor)\n│   ├── Verify OTP (Vendor) ← Tests script saves vendor_token\n│   ├── Send OTP (Customer)\n│   ├── Verify OTP (Customer) ← Tests script saves customer_token + customer_id\n│   ├── Get Profile\n│   ├── Update Profile\n│   ├── Update Location\n│   ├── Refresh Token\n│   └── Logout\n├── 2. Stores\n│   ├── Create Store ← Tests script saves store_id\n│   ├── Nearby Stores\n│   ├── Store Detail\n│   ├── Update Store\n│   ├── Follow Store\n│   ├── Review Store\n│   └── QR Code\n├── 3. Products\n│   ├── Create Product ← Tests script saves product_id\n│   ├── Nearby Products\n│   ├── Search Products\n│   ├── Product Detail\n│   ├── Update Product\n│   ├── Delete Product\n│   └── Wishlist Toggle\n├── 4. Videos\n│   ├── Request Upload URL ← Tests script saves video_id\n│   ├── Confirm Upload\n│   ├── My Videos (Vendor Library)\n│   ├── Update Video\n│   ├── Video Feed\n│   ├── Video Detail\n│   ├── Like / Unlike Video\n│   └── Delete Video\n├── 5. Chat\n│   ├── Start Conversation ← Tests script saves conversation_id\n│   ├── List Conversations (Inbox)\n│   ├── Message History\n│   ├── Mark as Read\n│   └── [WebSocket] Connect + Send  (use Postman WS tab or wscat)\n├── 6. Blacklist\n│   ├── Block Customer  (vendor token + store_id + customer_id)\n│   ├── List Blocked Customers  (vendor token + store_id)\n│   └── Unblock Customer  (same endpoint as Block — toggle)\n├── 7. Billing\n│   ├── List Plans  (no auth required)\n│   ├── Wallet Balance  (vendor token)\n│   ├── Top Up Wallet  (vendor token + {"amount": "1000.00"})\n│   ├── Subscribe to Plan  (vendor token + {"plan_name": "basic"})\n│   ├── Subscription Status  (vendor token)\n│   └── Transaction History  (vendor token)\n├── 8. Analytics + Admin\n│   ├── [Analytics] Vendor Dashboard  (vendor token)\n│   ├── [Analytics] Vendor Video Stats  (vendor token)\n│   ├── [Analytics] Vendor Product Stats  (vendor token)\n│   ├── [Admin] Platform Stats  (admin_token)\n│   ├── [Admin] List All Stores  (admin_token)\n│   ├── [Admin] Verify Store  (admin_token + PATCH {"is_verified": true})\n│   ├── [Admin] List All Users  (admin_token)\n│   └── [Admin] Toggle User Active  (admin_token + POST empty body)\n├── 9. Reservations\n│   ├── Create Reservation  (customer_token + store_id + product_id)\n│   ├── List Reservations  (customer_token → own / vendor_token → store\'s)\n│   ├── Reservation Detail  (customer_token or vendor_token)\n│   ├── Update Status — Confirm  (vendor_token + {"status": "confirmed"})\n│   ├── Update Status — Complete  (vendor_token + {"status": "completed"})\n│   ├── Update Status — Cancel  (vendor_token + {"status": "cancelled"})\n│   └── Customer Cancel  (customer_token + POST empty body)\n└── 10. Groups\n    ├── Get My Profile ID  (any_token → GET /auth/me/)\n    ├── Search by Profile ID  (any_token + ?profile_id=NK-XXXXXX)\n    ├── Create Customer Group  ← Tests saves group_id\n    ├── Create Vendor Group  (vendor_token)\n    ├── List My Groups\n    ├── Group Detail  (member token)\n    ├── Add Member — by Profile ID  (admin_token + {"profile_id": "NK-..."})\n    ├── Eligible Members  (vendor admin_token — followers not in group)\n    ├── Add Member — by User ID  (vendor admin_token + {"user_id": "..."})\n    ├── Make Admin  (admin_token + POST empty body)\n    ├── Remove Admin  (admin_token + POST empty body)\n    ├── Share Product  ← Tests saves sp_id\n    ├── Share Product — External Link  (should return 400)\n    ├── List Shared Products\n    ├── Finalize Product  (admin_token + POST empty body)\n    ├── Leave Group  (member_token)\n    └── Delete Group  (creator_token)\n└── 11. Notifications\n    ├── Register Device Token  (any_token + {"fcm_token": "...", "device_type": "android"})\n    ├── List Notifications (Inbox)  ← Tests saves notification_id\n    ├── Unread Count (Badge)\n    ├── Mark One Read  (any_token + notification_id)\n    └── Mark All Read  (any_token)\n└── 12. Store Hours + Razorpay\n    ├── Get Store Hours  (vendor_token + store_id)\n    ├── Set Store Hours  (vendor_token + store_id + array body) ← replaces all hours\n    ├── [Payment] List Plans  (no auth — check prices before initiating)\n    ├── [Payment] Initiate Payment  (vendor_token + {"plan_name": "basic"}) ← Tests saves razorpay_order_id\n    ├── [Payment] Verify Payment  (vendor_token + 4 fields) ← subscription activated\n    ├── [Payment] Webhook Simulation  (no auth + X-Razorpay-Signature header)\n    ├── [Payment] Webhook Duplicate  (same body again → already_processed)\n    ├── Check Subscription After Payment  (vendor_token → GET /billing/subscription/)\n    └── Check Transactions After Payment  (vendor_token → GET /billing/transactions/)')

h2('Auto-Save Token Script  (paste in Tests tab of Verify OTP request)')
code('// For VENDOR verify OTP request:\nconst r = pm.response.json();\nif (r.access) {\n    pm.environment.set("vendor_token", r.access);\n    pm.environment.set("vendor_refresh", r.refresh);\n    console.log("Vendor token saved:", r.access.substring(0, 30) + "...");\n}\n\n// For CUSTOMER verify OTP request — use this instead:\n// pm.environment.set("customer_token", r.access);\n// pm.environment.set("customer_refresh", r.refresh);')

h2('Auto-Save Store ID  (paste in Tests tab of Create Store request)')
code('const r = pm.response.json();\nif (r.id) {\n    pm.environment.set("store_id", r.id);\n    console.log("store_id saved:", r.id);\n}')

h2('Auto-Save Product ID  (paste in Tests tab of Create Product request)')
code('const r = pm.response.json();\nif (r.id) {\n    pm.environment.set("product_id", r.id);\n    console.log("product_id saved:", r.id);\n}')

h2('Auto-Save Video ID  (paste in Tests tab of Request Upload URL request)')
code('const r = pm.response.json();\nif (r.video_id) {\n    pm.environment.set("video_id", r.video_id);\n    pm.environment.set("upload_url", r.upload_url);\n    console.log("video_id saved:", r.video_id);\n}')

h2('Auto-Save Conversation ID  (paste in Tests tab of Start Conversation request)')
code('const r = pm.response.json();\nif (r.id) {\n    pm.environment.set("conversation_id", r.id);\n    console.log("conversation_id saved:", r.id);\n}')

h2('Auto-Save Customer ID  (paste in Tests tab of Verify OTP Customer request)')
code('// Add this to the existing customer verify OTP Tests script:\nconst r = pm.response.json();\nif (r.user && r.user.id) {\n    pm.environment.set("customer_id", r.user.id);\n    console.log("customer_id saved:", r.user.id);\n}')
tip('customer_id is needed for Sprint 6 blacklist endpoint:\n'
    '  POST /stores/{{store_id}}/blacklist/{{customer_id}}/')

h2('Quick Test — All Public Endpoints via curl')
para('Copy-paste these into terminal to test without Postman:')
code('# Health check\ncurl -s http://localhost:8000/api/v1/health/ | python3 -m json.tool\n\n# Send OTP\ncurl -s -X POST http://localhost:8000/api/v1/auth/otp/send/ \\\n  -H "Content-Type: application/json" \\\n  -d \'{"phone_number": "+919999999999"}\'\n\n# Verify OTP (copy access token from response)\ncurl -s -X POST http://localhost:8000/api/v1/auth/otp/verify/ \\\n  -H "Content-Type: application/json" \\\n  -d \'{"phone_number": "+919999999999", "otp": "123456"}\'\n\n# Get profile (replace TOKEN with your access token)\ncurl -s http://localhost:8000/api/v1/auth/me/ \\\n  -H "Authorization: Bearer TOKEN"')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 6 — FUTURE SPRINT UPDATE GUIDE
# ════════════════════════════════════════════════════════════════
h1('6. How to Update This Document (Future Sprints)')

note('📌  IMPORTANT: When a new sprint adds new API endpoints, update this document.\n'
     '     Do NOT create a separate doc — keep everything in this single file.',
     'D1ECF1', (0x0C, 0x54, 0x60))

doc.add_paragraph()
h2('Checklist for each new sprint:')
bullet('Add new endpoint rows to Section 2 — Complete API Reference table')
bullet('Add new STEP sections in Section 3 — Step-by-Step Test Flow')
bullet('Add new error cases in Section 4 — Error Cases to Test')
bullet('Add new Postman folder in Section 5 — Postman Quick Setup')
bullet('Update the cover page "Last updated" date and sprint range')
bullet('If new field types are introduced (e.g. file uploads), add a format rules subsection in Section 1')

h2('Upcoming Sprints to Add Here')
table(
    ['Sprint', 'Module', 'APIs to Add'],
    [
        ('Sprint 4', 'Video Module',        '✅ DONE — Steps 24–31 in this document'),
        ('Sprint 5', 'Chat Module',         '✅ DONE — Steps 32–36 in this document'),
        ('Sprint 6', 'Blacklist Engine',    '✅ DONE — Steps 37–38 in this document'),
        ('Sprint 7', 'Billing / Wallet',    '✅ DONE — Steps 39–44 in this document'),
        ('Sprint 8', 'Analytics + Admin',   '✅ DONE — Steps 45–53 in this document'),
        ('Sprint 9',  'Reservations',        '✅ DONE — Steps 54–58 in this document'),
        ('Sprint 10', 'Groups',             '✅ DONE — Steps 59–69 in this document'),
        ('Sprint 11', 'Notifications',      '✅ DONE — Steps 70–74 in this document'),
        ('Sprint 12', 'Production + Razorpay', '✅ DONE — Steps 75–79 in this document'),
    ],
    col_widths=[0.8, 1.5, 4.2]
)

# ── FOOTER ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('NearKart Backend — Master API Testing Guide   |   Sprint 1–12   |   Use Postman or curl for testing')
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── SAVE ──
out = '/Users/hazeevali/Documents/NearKart/Backend/nearkart_backend/docs/NearKart_Master_API_Testing_Guide.docx'
doc.save(out)
print(f'Saved: {out}')
